#!/usr/bin/env python3
# generate_dokumen_ppk.py — Generate dokumen PPK PL dari Master Data Excel
# Mode: generate | pdf | paket_baru | multi | commit-paket | nomor-baru | test-terbilang
import sys, os, shutil, re as _re, zipfile, argparse
from datetime import date, datetime, timedelta

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

RUNTIME_DIR = os.path.dirname(os.path.abspath(__file__))


def _configured_path(name: str) -> str:
    value = os.environ.get(name, '').strip().strip('"')
    return os.path.normpath(value) if value else ''


# Runtime berada di Google Drive bersama workbook, sedangkan source dan Python
# berada di clone lokal per-PC. Launcher mengirim PPK_PACKAGE_ROOT saat dipanggil
# dari VBA; fallback relatif mempertahankan kompatibilitas paket lama.
BASE_DIR = _configured_path('PPK_PACKAGE_ROOT') or os.path.normpath(os.path.join(RUNTIME_DIR, '..'))
_EXCEL_CANDIDATES = [
    os.path.join(BASE_DIR, '0. Master_Data_PL_PPK.xlsm'),
    os.path.join(BASE_DIR, '0. Master_Data_PL_PPK (1).xlsm'),
]
EXCEL_PATH = next((p for p in _EXCEL_CANDIDATES if os.path.isfile(p)), _EXCEL_CANDIDATES[0])
LOG_DIR    = _configured_path('PPK_LOG_DIR') or RUNTIME_DIR
LOG_PATH   = os.path.join(LOG_DIR, 'generate_log.txt')
LOG_XLSX   = os.path.join(LOG_DIR, 'ppk_log.xlsx')

# Resolver source/runtime lokal per-PC; tidak pernah mencari source di Google Drive.
_V19_ROOT = _configured_path('POKJA_V19_ROOT')
_CODE_ROOT = _configured_path('POKJA_CODE_ROOT')
_REPO_ROOT = _V19_ROOT or _CODE_ROOT
_CANDIDATES = [
    os.path.join(_V19_ROOT, 'python', 'Lib', 'site-packages') if _V19_ROOT else '',
    os.path.join(_CODE_ROOT, '..', 'Runtime', 'WPy64-313110', 'python', 'Lib', 'site-packages') if _CODE_ROOT else '',
    os.path.join(os.path.dirname(sys.executable), 'Lib', 'site-packages'),
]
for _p in _CANDIDATES:
    if _p and os.path.isdir(_p) and _p not in sys.path:
        sys.path.insert(0, _p)
        break

# Secret hanya lokal per-PC. Generator biasa tidak membutuhkan secret; mode
# muat-db akan memberi error eksplisit bila secret belum disiapkan.
_SECRET_ROOT = _configured_path('POKJA_SECRET_ROOT') or (
    os.path.join(os.environ.get('LOCALAPPDATA', ''), 'POKJA2026', 'Secrets')
    if os.environ.get('LOCALAPPDATA') else ''
)
_SB_CANDIDATES = [
    os.path.join(_SECRET_ROOT, 'secret_supabase.env') if _SECRET_ROOT else '',
    os.path.join(_V19_ROOT, 'secret_supabase.env') if _V19_ROOT else '',
]
SB_ENV = next((p for p in _SB_CANDIDATES if p and os.path.exists(p)), _SB_CANDIDATES[0] if _SB_CANDIDATES else '')

# ─── Field wajib (validasi sebelum generate) ──────────────────────────────────
WAJIB = [
    "Nama Paket (Lengkap)", "Kode RUP",
    "Pagu Anggaran (Angka)", "Nilai HPS (Angka)", "Jangka Waktu (Hari)",
    "Lokasi Pekerjaan", "Tahun Anggaran",
    "Nama PPK", "NIP PPK", "Jabatan PPK",
    "Tanggal KAK & HPS",
]

# ─── Field yang direset saat "Paket Baru" (Seksi B & C tidak direset) ─────────
RESET_FIELDS = [
    "Nama Paket (Lengkap)", "Kode RUP",
    "Kode Jenis Paket",
    "Kode Rekening (MAK)", "Metode Pemilihan", "Jenis Kontrak",
    "Tahap Dokumen",
    "Pagu Anggaran (Angka)", "Pagu Anggaran (Terbilang)",
    "Nilai HPS (Angka)", "Nilai HPS (Terbilang)",
    "Jangka Waktu (Hari)", "Jangka Waktu (Terbilang)",
    "Lokasi Pekerjaan", "Sumber Dana", "Program", "Kegiatan",
    "Sub Kegiatan", "SBU yang Disyaratkan", "Nomor DIPA/DPA",
    # Seksi D
    "Tanggal KAK & HPS", "Bulan Tahun KAK & HPS",
    "Tanggal Nota Dinas", "Nomor Nota Dinas",
    "Tanggal SPPBJ", "Nomor SPPBJ",
    "Tanggal SPK", "Nomor SPK",
    "Tanggal SPMK", "Nomor SPMK",
    "Tanggal Mulai Kerja", "Tanggal Selesai Kerja",
    # Seksi E
    "Nama Penyedia", "Alamat Penyedia", "Nama Direktur", "Jabatan Direktur",
    "Nomor Akta Penyedia", "Tanggal Akta Penyedia", "Nama Notaris Penyedia", "NPWP Penyedia",
    "Nama Bank", "Nomor Rekening", "Atas Nama Rekening",
    "Nomor Surat Minat", "Tanggal Surat Minat",
    "Nilai Penawaran (Angka)", "Nilai Penawaran (Terbilang)",
    "Nilai Kontrak (Angka)", "Nilai Kontrak (Terbilang)",
    # Seksi F
    "Nama Ketua Tim", "SKA/SKK Ketua Tim", "Pendidikan Ketua Tim",
    "Nama Surveyor", "Nama Asisten Surveyor", "Nama Estimator/Drafter",
    # Seksi G
    "Nomor Surat Undangan PL", "Tanggal Surat Undangan PL",
    "Nomor BA Hasil PL", "Tanggal BA Hasil PL",
    # Seksi H: substansi kontrak PK
    "Uang Muka (%)", "Sistem Pembayaran", "Rincian Termin",
    "Masa Pemeliharaan (Hari)", "Retensi (%)",
    "Ada Wakil Sah?", "Nama Wakil Sah", "Jabatan Wakil Sah",
    "NIP Wakil Sah", "Nomor SK Wakil Sah", "Tanggal SK Wakil Sah",
]

# ─── FIELD_MAP: label Excel → placeholder Word ────────────────────────────────
FIELD_MAP = {
    "Nama Paket (Lengkap)":         "«NAMA_PAKET_LENGKAP»",
    "Kode RUP":                     "«KODE_RUP»",
    "Pagu Anggaran (Angka)":        "«PAGU_ANGKA»",
    "Pagu Anggaran (Terbilang)":    "«PAGU_TERBILANG»",
    "Nilai HPS (Angka)":            "«HPS_ANGKA»",
    "Nilai HPS (Terbilang)":        "«HPS_TERBILANG»",
    "Jangka Waktu (Hari)":          "«JANGKA_WAKTU_HARI»",
    "Jangka Waktu (Terbilang)":     "«JANGKA_WAKTU_TERBILANG»",
    "Lokasi Pekerjaan":             "«LOKASI_PEKERJAAN»",
    "Sumber Dana":                  "«SUMBER_DANA»",
    "Tahun Anggaran":               "«TAHUN_ANGGARAN»",
    "Program":                      "«PROGRAM»",
    "Kegiatan":                     "«KEGIATAN»",
    "Sub Kegiatan":                 "«SUB_KEGIATAN»",
    "Nomor DIPA/DPA":               "«NOMOR_DIPA_DPA»",
    "Lingkup Pekerjaan":            "«LINGKUP_PEKERJAAN»",
    "SBU yang Disyaratkan":         "«SBU»",
    "Kota Dokumen":                 "«KOTA_DOKUMEN»",
    "Nama SKPD/OPD":                "«NAMA_SKPD»",
    "Nama SKPD Singkat":            "«NAMA_SKPD_SINGKAT»",
    "Alamat SKPD":                  "«ALAMAT_SKPD»",
    "Kabupaten/Kota":               "«KABUPATEN_KOTA»",
    "Nama PPK":                     "«NAMA_PPK»",
    "NIP PPK":                      "«NIP_PPK»",
    "Nama PA":                    "«NAMA_PA»",
    "NIP PA":                     "«NIP_PA»",
    "Nomor Surat Diskresi":       "«NOMOR_SURAT_DISKRESI»",
    "Tanggal Diskresi":           "«TANGGAL_DISKRESI»",
    "Jabatan PPK":                  "«JABATAN_PPK»",
    "Nomor SK PPK":                 "«NOMOR_SK_PPK»",
    "Tanggal SK PPK":               "«TANGGAL_SK_PPK»",
    "Uraian SK PPK":                "«URAIAN_SK_PPK»",
    "Tanggal KAK & HPS":            "«TANGGAL_KAK_HPS»",
    "Bulan Tahun KAK & HPS":        "«BULAN_TAHUN_KAK_HPS»",
    "Tanggal Nota Dinas":           "«TANGGAL_NOTA_DINAS»",
    "Nomor Nota Dinas":             "«NOMOR_NOTA_DINAS»",
    "Tanggal SPPBJ":                "«TANGGAL_SPPBJ»",
    "Nomor SPPBJ":                  "«NOMOR_SPPBJ»",
    "Tanggal SPK":                  "«TANGGAL_SPK»",
    "Nomor SPK":                    "«NOMOR_SPK»",
    "Tanggal SPMK":                 "«TANGGAL_SPMK»",
    "Nomor SPMK":                   "«NOMOR_SPMK»",
    "Tanggal Mulai Kerja":          "«TANGGAL_MULAI_KERJA»",
    "Tanggal Selesai Kerja":        "«TANGGAL_SELESAI_KERJA»",
    "Nama Penyedia":                "«NAMA_PENYEDIA»",
    "Alamat Penyedia":              "«ALAMAT_PENYEDIA»",
    "Nama Direktur":                "«NAMA_DIREKTUR»",
    "Jabatan Direktur":              "«JABATAN_DIREKTUR»",
    "NPWP Penyedia":                "«NPWP_PENYEDIA»",
    "Nama Bank":                    "«NAMA_BANK»",
    "Nomor Rekening":               "«NOMOR_REKENING»",
    "Atas Nama Rekening":           "«ATAS_NAMA_REKENING»",
    "Nomor Surat Minat":            "«NOMOR_SURAT_MINAT»",
    "Tanggal Surat Minat":          "«TANGGAL_SURAT_MINAT»",
    "Nilai Penawaran (Angka)":      "«NILAI_PENAWARAN_ANGKA»",
    "Nilai Penawaran (Terbilang)":  "«NILAI_PENAWARAN_TERBILANG»",
    "Nilai Kontrak (Angka)":        "«NILAI_KONTRAK_ANGKA»",
    "Nilai Kontrak (Terbilang)":    "«NILAI_KONTRAK_TERBILANG»",
    "Nama Ketua Tim":               "«NAMA_KETUA_TIM»",
    "SKA/SKK Ketua Tim":            "«SKA_KETUA_TIM»",
    "Pendidikan Ketua Tim":         "«PENDIDIKAN_KETUA_TIM»",
    "Nama Surveyor":                "«NAMA_SURVEYOR»",
    "Nama Asisten Surveyor":        "«NAMA_ASISTEN_SURVEYOR»",
    "Nama Estimator/Drafter":       "«NAMA_ESTIMATOR»",
    "Nomor Surat Undangan PL":      "«NOMOR_UNDANGAN_PL»",
    "Tanggal Surat Undangan PL":    "«TANGGAL_UNDANGAN_PL»",
    "Nomor BA Hasil PL":            "«NOMOR_BA_PL»",
    "Tanggal BA Hasil PL":          "«TANGGAL_BA_PL»",
}


# Additional stage-aware contract fields. Kept outside the legacy map block so
# old placeholder mappings remain untouched.
FIELD_MAP.update({
    "Tahap Dokumen": "\u00abTAHAP_DOKUMEN\u00bb",
    "Ada Wakil Sah?": "\u00abADA_WAKIL_SAH\u00bb",
    "Nama Wakil Sah": "\u00abNAMA_WAKIL_SAH\u00bb",
    "Jabatan Wakil Sah": "\u00abJABATAN_WAKIL_SAH\u00bb",
    "NIP Wakil Sah": "\u00abNIP_WAKIL_SAH\u00bb",
    "Nomor SK Wakil Sah": "\u00abNOMOR_SK_WAKIL_SAH\u00bb",
    "Tanggal SK Wakil Sah": "\u00abTANGGAL_SK_WAKIL_SAH\u00bb",
    "Uang Muka (%)": "\u00abUANG_MUKA_PERSEN\u00bb",
    "Sistem Pembayaran": "\u00abSISTEM_PEMBAYARAN\u00bb",
    "Rincian Termin": "\u00abRINCIAN_TERMIN\u00bb",
    "Masa Pemeliharaan (Hari)": "\u00abMASA_PEMELIHARAAN_HARI\u00bb",
    "Retensi (%)": "\u00abRETENSI_PERSEN\u00bb",
    "Jenis Kontrak": "\u00abJENIS_KONTRAK\u00bb",
    "Nomor Akta Penyedia": "\u00abNOMOR_AKTA_PENYEDIA\u00bb",
    "Tanggal Akta Penyedia": "\u00abTANGGAL_AKTA_PENYEDIA\u00bb",
    "Nama Notaris Penyedia": "\u00abNAMA_NOTARIS_PENYEDIA\u00bb",
})

DOCUMENT_STAGES = {
    "UPLOAD AWAL": 0,
    "BERKONTRAK": 3,
}

STAGE_REQUIRED_FIELDS = {
    "BERKONTRAK": [
        "Nama Penyedia", "Alamat Penyedia", "Nomor SPPBJ", "Tanggal SPPBJ",
        "Nilai Penawaran (Angka)", "Nilai Penawaran (Terbilang)",
        "Nama Direktur", "Jabatan Direktur",
        "Nama Bank", "Nomor Rekening", "Atas Nama Rekening",
        "Nomor Akta Penyedia", "Tanggal Akta Penyedia", "Nama Notaris Penyedia",
        "Nomor SPK", "Tanggal SPK", "Nilai Kontrak (Angka)",
        "Nilai Kontrak (Terbilang)", "Uang Muka (%)", "Sistem Pembayaran",
        "Nomor SPMK", "Tanggal SPMK",
        "Tanggal Mulai Kerja", "Tanggal Selesai Kerja",
    ],
}

LEGACY_FINAL_STAGES = {
    "SPPBJ FINAL": "BERKONTRAK",
    "SPK FINAL": "BERKONTRAK",
    "SPMK FINAL": "BERKONTRAK",
}

CONTRACT_FINAL_FIELDS = set(STAGE_REQUIRED_FIELDS["BERKONTRAK"])

DOCUMENT_STAGE_LABELS = {
    "UPLOAD AWAL": "UPLOAD AWAL",
    "BERKONTRAK": "BERKONTRAK",
}

CONTRACT_FIELD_KEYS = {
    "Tanggal SPPBJ", "Nomor SPPBJ", "Tanggal SPK", "Nomor SPK",
    "Tanggal SPMK", "Nomor SPMK", "Tanggal Mulai Kerja",
    "Tanggal Selesai Kerja", "Nama Penyedia", "Alamat Penyedia",
    "Nama Direktur", "Jabatan Direktur", "Nomor Akta Penyedia", "Tanggal Akta Penyedia",
    "Nama Notaris Penyedia", "NPWP Penyedia", "Nama Bank", "Nomor Rekening",
    "Atas Nama Rekening", "Nilai Penawaran (Angka)",
    "Nilai Penawaran (Terbilang)", "Nilai Kontrak (Angka)",
    "Nilai Kontrak (Terbilang)", "Nomor Surat Undangan PL",
    "Tanggal Surat Undangan PL", "Nomor BA Hasil PL", "Tanggal BA Hasil PL",
    "Uang Muka (%)", "Sistem Pembayaran", "Rincian Termin",
    "Masa Pemeliharaan (Hari)", "Retensi (%)",
    "Ada Wakil Sah?", "Nama Wakil Sah", "Jabatan Wakil Sah",
    "NIP Wakil Sah", "Nomor SK Wakil Sah", "Tanggal SK Wakil Sah",
}

def normalize_document_stage(value: object) -> str:
    stage = str(value or "UPLOAD AWAL").strip().upper()
    stage = LEGACY_FINAL_STAGES.get(stage, stage)
    return stage if stage in DOCUMENT_STAGES else "UPLOAD AWAL"


def contract_doc_kind(filename: str) -> str | None:
    name = os.path.basename(filename).lower()
    if name.startswith("3. sppbj"):
        return "sppbj"
    if name.startswith("4. spmk"):
        return "spmk"
    if name.startswith("5. ") and "spk" in name:
        return "spk"
    return None


def detail_enabled(stage: str, kind: str | None) -> bool:
    if kind is None:
        return True
    level = DOCUMENT_STAGES.get(stage, 0)
    required_level = {"sppbj": 1, "spk": 2, "spmk": 3}[kind]
    return level >= required_level


def _percentage_value(value):
    if value is None or str(value).strip() in ("", "None"):
        return None
    try:
        number = float(str(value).strip().replace("%", "").replace(",", "."))
        return number * 100 if 0 <= number <= 1 else number
    except (TypeError, ValueError):
        return None


def _minimum_uang_muka(excel_data: dict) -> float:
    hps = _numeric_value(excel_data.get("Nilai HPS (Angka)"))
    return 50.0 if hps <= 200_000_000 else 30.0


def _format_percentage(value) -> str:
    number = _percentage_value(value)
    return f"{number:g}%" if number is not None else str(value)


_BULAN_ID = {
    "januari": 1, "februari": 2, "pebruari": 2,
    "maret": 3, "april": 4, "mei": 5, "juni": 6,
    "juli": 7, "agustus": 8, "september": 9, "oktober": 10,
    "november": 11, "desember": 12,
}


def _parse_tanggal(value):
    """Parse tanggal Excel yang bisa berupa 01082026 atau teks Indonesia."""
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    text = str(value or "").strip()
    if not text or text.lower() == "none":
        return None
    digits = _re.sub(r"\D", "", text)
    if len(digits) in (7, 8) and len(digits) <= 8:
        if len(digits) == 7:
            digits = "0" + digits
        try:
            return date(int(digits[4:]), int(digits[2:4]), int(digits[:2]))
        except ValueError:
            return None
    match = _re.match(r"^(\d{1,2})[\s./-]+([A-Za-z]+)[\s./-]+(\d{4})$", text.lower())
    if match and match.group(2) in _BULAN_ID:
        try:
            return date(int(match.group(3)), _BULAN_ID[match.group(2)], int(match.group(1)))
        except ValueError:
            return None
    match = _re.match(r"^(\d{1,2})[./-](\d{1,2})[./-](\d{4})$", text)
    if match:
        try:
            return date(int(match.group(3)), int(match.group(2)), int(match.group(1)))
        except ValueError:
            return None
    return None


def _format_tanggal(value) -> str:
    parsed = _parse_tanggal(value)
    if not parsed:
        return str(value or "")
    bulan = (
        "Januari", "Februari", "Maret", "April", "Mei", "Juni",
        "Juli", "Agustus", "September", "Oktober", "November", "Desember",
    )
    return f"{parsed.day} {bulan[parsed.month - 1]} {parsed.year}"


def _auto_tanggal_selesai(excel_data: dict) -> dict:
    """Isi tanggal mulai fallback SPK dan tanggal selesai secara inklusif."""
    if normalize_document_stage(excel_data.get("Tahap Dokumen")) != "BERKONTRAK":
        return excel_data
    start = _parse_tanggal(excel_data.get("Tanggal Mulai Kerja"))
    if start is None:
        start = _parse_tanggal(excel_data.get("Tanggal SPK"))
        if start is not None:
            excel_data["Tanggal Mulai Kerja"] = _format_tanggal(start)
    days = _numeric_value(excel_data.get("Jangka Waktu (Hari)"))
    if start is not None and days > 0:
        excel_data["Tanggal Selesai Kerja"] = _format_tanggal(
            start + timedelta(days=max(0, int(days) - 1))
        )
    return excel_data


def _source_funds_detail(excel_data: dict) -> str:
    source = str(excel_data.get("Sumber Dana", "") or "").strip()
    skpd = str(excel_data.get("Nama SKPD/OPD", "") or "").strip()
    dipa = str(excel_data.get("Nomor DIPA/DPA", "") or "").strip()
    year = str(excel_data.get("Tahun Anggaran", "") or "").strip()
    mak = str(excel_data.get("Kode Rekening (MAK)", "") or "").strip()
    parts = [source] if source else []
    if skpd:
        parts.append(f"dibebankan atas DIPA/DPA {skpd}")
    elif dipa:
        # Fallback workbook lama yang belum memiliki Nama SKPD/OPD.
        parts.append(f"dibebankan atas DIPA/DPA {dipa}")
    if year:
        parts.append(f"Tahun Anggaran {year}")
    mata_anggaran = dipa or mak
    if mata_anggaran:
        parts.append(f"untuk mata anggaran kegiatan {mata_anggaran}")
    return ", ".join(parts)


def validate_document_stage(excel_data: dict) -> list[str]:
    stage = normalize_document_stage(excel_data.get("Tahap Dokumen"))
    required = STAGE_REQUIRED_FIELDS.get(stage, [])
    missing = [
        key for key in required
        if key != "Uang Muka (%)"
        and not str(excel_data.get(key, "") or "").strip()
    ]
    payment = str(excel_data.get("Sistem Pembayaran", "") or "").strip().upper()
    if stage == "BERKONTRAK" and payment == "TERMIN" and not str(excel_data.get("Rincian Termin", "") or "").strip():
        missing.append("Rincian Termin")
    if stage == "BERKONTRAK":
        uang_muka = _percentage_value(excel_data.get("Uang Muka (%)"))
        if uang_muka is not None and uang_muka < _minimum_uang_muka(excel_data):
            missing.append(
                f"Uang Muka (%) minimal {_minimum_uang_muka(excel_data):g}% "
                f"untuk HPS {_numeric_value(excel_data.get('Nilai HPS (Angka)')):,.0f}"
            )
        for key in ("Tanggal SPK", "Tanggal SPMK", "Tanggal Mulai Kerja", "Tanggal Selesai Kerja"):
            if str(excel_data.get(key, "") or "").strip() and _parse_tanggal(excel_data.get(key)) is None:
                missing.append(f"{key} (format tanggal tidak dikenali)")
        wakil = str(excel_data.get("Ada Wakil Sah?", "") or "").strip().upper()
        if wakil not in ("", "ADA", "TIDAK ADA", "TIDAK"):
            missing.append("Ada Wakil Sah? (isi ADA atau TIDAK ADA)")
        if wakil == "ADA":
            for key in (
                "Nama Wakil Sah", "Jabatan Wakil Sah", "NIP Wakil Sah",
                "Nomor SK Wakil Sah", "Tanggal SK Wakil Sah",
            ):
                if not str(excel_data.get(key, "") or "").strip():
                    missing.append(key)
            if str(excel_data.get("Tanggal SK Wakil Sah", "") or "").strip() and _parse_tanggal(excel_data.get("Tanggal SK Wakil Sah")) is None:
                missing.append("Tanggal SK Wakil Sah (format tanggal tidak dikenali)")
    return missing


# ═══════════════════════════════════════════════════════════════════════════════
# F1: TERBILANG — pure Python, tanpa library eksternal
# ═══════════════════════════════════════════════════════════════════════════════
_SATUAN = ['', 'Satu', 'Dua', 'Tiga', 'Empat', 'Lima',
           'Enam', 'Tujuh', 'Delapan', 'Sembilan']
_BELASAN = {
    11: 'Sebelas', 12: 'Dua Belas', 13: 'Tiga Belas', 14: 'Empat Belas',
    15: 'Lima Belas', 16: 'Enam Belas', 17: 'Tujuh Belas',
    18: 'Delapan Belas', 19: 'Sembilan Belas',
}


def _terbilang_ratus(n: int) -> str:
    if n == 0:
        return ''
    if n < 10:
        return _SATUAN[n]
    if 11 <= n <= 19:
        return _BELASAN[n]
    if n == 10:
        return 'Sepuluh'
    if n < 100:
        puluhan = n // 10
        satuan  = n % 10
        prefix  = 'Sepuluh' if puluhan == 1 else _SATUAN[puluhan] + ' Puluh'
        return (prefix + ' ' + _SATUAN[satuan]).strip()
    ratus  = n // 100
    sisa   = n % 100
    prefix = 'Seratus' if ratus == 1 else _SATUAN[ratus] + ' Ratus'
    return (prefix + ' ' + _terbilang_ratus(sisa)).strip()


def terbilang_angka(angka) -> str:
    """Konversi angka ke teks Bahasa Indonesia (tanpa satuan)."""
    try:
        n = int(str(angka).replace('.', '').replace(',', '').strip())
    except (ValueError, TypeError):
        return str(angka)
    if n == 0:
        return 'Nol'
    if n < 0:
        return 'Minus ' + terbilang_angka(-n)
    parts = []
    for nilai, nama in [(10**12, 'Triliun'), (10**9, 'Miliar'),
                        (10**6, 'Juta'), (10**3, 'Ribu'), (1, '')]:
        if n >= nilai:
            kelompok = n // nilai
            n        = n % nilai
            if nilai == 1000 and kelompok == 1:
                parts.append('Seribu')
            else:
                teks = _terbilang_ratus(kelompok)
                parts.append((teks + ' ' + nama).strip() if nama else teks)
    return ' '.join(p for p in parts if p)


def terbilang_rupiah(angka) -> str:
    return terbilang_angka(angka) + ' Rupiah'


def terbilang_hari(n) -> str:
    try:
        return terbilang_angka(int(str(n).strip()))
    except Exception:
        return str(n)


# ═══════════════════════════════════════════════════════════════════════════════
# F4: TEMPLATE AUTO-DETECT
# ═══════════════════════════════════════════════════════════════════════════════
def template_auto_detect(base_dir: str):
    """Scan .docx di base_dir yang berawalan angka atau berisi placeholder «...»."""
    result = []
    def _num_key(name):
        m2 = _re.match(r'^(\d+)', name)
        return (int(m2.group(1)) if m2 else 9999, name)
    for f in sorted(os.listdir(base_dir), key=_num_key):
        if not f.lower().endswith('.docx') or f.startswith('~$'):
            continue
        # Pola deteksi prefix angka: "1.", "2.", dst.
        m = _re.match(r'^(\d+\.[^.]+?)(?:\s+\w.*)?\.docx$', f, _re.IGNORECASE)
        is_numbered = bool(m)

        path = os.path.join(base_dir, f)
        try:
            with zipfile.ZipFile(path) as z:
                content = z.read('word/document.xml').decode('utf-8', errors='ignore')
                # Masukkan jika berawalan angka ATAU mengandung placeholder '«'
                if is_numbered or '«' in content:
                    prefix = m.group(1).strip() if m else None
                    result.append((f, prefix))
        except Exception:
            if is_numbered:  # Tetap masukkan jika berawalan angka meski zip corrupt/gagal baca
                prefix = m.group(1).strip() if m else None
                result.append((f, prefix))
    return result


def _word_files_fallback():
    return [
        ("1. KAK Pembuatan Pagar Pasar Rabu Suato Tatakan.docx",  "1. KAK"),
        ("2. Uraian Singkat Pekerjaan Pagar Pasar Rabu.docx",     "2. Uraian Singkat Pekerjaan"),
        ("3. SPPBJ.docx",                                          "3. SPPBJ"),
        ("4. SPMK.docx",                                           "4. SPMK"),
        ("5. Rancangan SPK.docx",                                  "5. Rancangan SPK"),
        ("6. SUK.docx",                                            None),
        ("7. HPS Pasar Rebo.docx",                                 "7. HPS"),
        ("8. Nota Dinas PPK Disdag.docx",                          "8. Nota Dinas PPK"),
        ("9. Daftar Personil Manajerial.docx",                     None),
    ]


TEMPLATE_PROFILES = {
    'PK': 'Konstruksi',
    'KP': 'Perencanaan',
    'KPWAS': 'Pengawasan',
}


def _template_profile(excel_data: dict) -> tuple[str, str]:
    """Resolve kode jenis paket dari Master Data ke folder template."""
    code = str(excel_data.get('Kode Jenis Paket', '') or '').strip().upper()
    if code not in TEMPLATE_PROFILES:
        raise ValueError(
            "Kode Jenis Paket wajib diisi PK, KP, atau KPWAS "
            f"(nilai saat ini: {code or 'kosong'})."
        )
    return code, TEMPLATE_PROFILES[code]


def _numeric_value(value):
    try:
        if isinstance(value, (int, float)):
            return float(value)
        return float(str(value).replace('.', '').replace(',', '.').strip())
    except (TypeError, ValueError):
        return 0.0


def template_files_for_data(base_dir: str, excel_data: dict):
    """Pilih template berdasarkan Kode Jenis Paket, bukan scan folder root."""
    code, profile_dir = _template_profile(excel_data)
    scan_dir = os.path.join(base_dir, profile_dir)
    if not os.path.isdir(scan_dir):
        raise FileNotFoundError(f"Folder template tidak ditemukan: {scan_dir}")

    pagu = _numeric_value(excel_data.get('Pagu Anggaran (Angka)'))
    result = []
    for name in sorted(os.listdir(scan_dir), key=lambda n: (int(_re.match(r'^\d+', n).group()) if _re.match(r'^\d+', n) else 9999, n)):
        lower_name = name.lower()
        if (
            not lower_name.endswith('.docx')
            or name.startswith('~$')
            or '.bak' in lower_name
            or '.before-' in lower_name
            or lower_name.endswith('.tmp.docx')
        ):
            continue
        # Konstruksi memiliki dua varian SPPBJ; pilih sesuai Pagu Anggaran.
        if code == 'PK' and name.lower().startswith('3. sppbj - sampai dengan 200 juta') and pagu > 200_000_000:
            continue
        if code == 'PK' and name.lower() == '3. sppbj.docx' and 0 < pagu <= 200_000_000:
            continue
        relative_name = os.path.join(profile_dir, name)
        result.append((relative_name, name[:-5]))
    if not result:
        raise ValueError(f"Tidak ada template Word aktif di folder {profile_dir}.")
    return result


def _personil_rows(excel_path: str, code: str):
    """Baca tabel personil sederhana dari Data PK; JKK legacy sebagai fallback."""
    from openpyxl import load_workbook
    wb = load_workbook(excel_path, read_only=True, data_only=True, keep_links=False)
    try:
        ws = wb['Data PK']
        start = 1 if code == 'PK' else 11
        rows = []
        for values in ws.iter_rows(min_row=5, max_row=19, min_col=start, max_col=start + 3, values_only=True):
            if not any(v not in (None, '') for v in values[1:]):
                continue
            rows.append({
                'No': values[0], 'Jabatan': values[1] or '',
                'Sertifikat': values[2] or '', 'Pengalaman Kerja': values[3] or '',
            })
        if rows or code == 'PK':
            return rows
        ws_legacy = wb['Data JKK']
        for values in ws_legacy.iter_rows(min_row=5, max_row=19, min_col=1, max_col=4, values_only=True):
            if any(v not in (None, '') for v in values[1:]):
                rows.append({'No': values[0], 'Jabatan': values[1] or '', 'Sertifikat': values[2] or '', 'Pengalaman Kerja': values[3] or ''})
        return rows
    finally:
        wb.close()


def merge_personil_template(template_path: str, excel_path: str, output_path: str, code: str):
    """Isi row marker List_Personil tanpa mengubah template asli."""
    import copy
    from lxml import etree
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns = {'w': W_NS}
    with zipfile.ZipFile(template_path, 'r') as zin:
        root = etree.fromstring(zin.read('word/document.xml'))
        target = None
        for table in root.xpath('.//w:tbl', namespaces=ns):
            text = ''.join(table.xpath('.//w:t/text()', namespaces=ns))
            if '[[JABATAN]]' in text:
                target = table
                break
        if target is None:
            shutil.copy2(template_path, output_path)
            return 0
        rows = target.xpath('./w:tr', namespaces=ns)
        if len(rows) < 2:
            raise ValueError(f'Tabel personil tidak memiliki row marker: {template_path}')
        template_row = rows[1]
        for extra in rows[2:]:
            target.remove(extra)
        people = _personil_rows(excel_path, code)
        for index, person in enumerate(people, 1):
            row = copy.deepcopy(template_row)
            mapping = {
                '[[NO]]': str(person.get('No') or index),
                '[[JABATAN]]': str(person.get('Jabatan', '')),
                '[[SERTIFIKAT]]': str(person.get('Sertifikat', '')),
                '[[PENGALAMAN]]': str(person.get('Pengalaman Kerja', '')),
            }
            for node in row.xpath('.//w:t', namespaces=ns):
                if node.text:
                    for old, new in mapping.items():
                        node.text = node.text.replace(old, new)
            target.append(row)
        target.remove(template_row)
        for node in root.xpath('.//w:t', namespaces=ns):
            if node.text:
                node.text = node.text.replace('[[DAFTAR_PERSONIL_JKK]]', '')
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            payload = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone='yes')
            for item in zin.infolist():
                zout.writestr(item, payload if item.filename == 'word/document.xml' else zin.read(item.filename))
    return len(people)


def _equipment_rows(excel_path: str):
    """Baca tabel peralatan sederhana PK dari kolom F:I sheet Data PK."""
    from openpyxl import load_workbook

    wb = load_workbook(excel_path, read_only=True, data_only=True, keep_links=False)
    try:
        ws = wb['Data PK']
        rows = []
        for values in ws.iter_rows(min_row=5, max_row=19, min_col=6, max_col=9, values_only=True):
            if not any(v not in (None, '') for v in values[1:]):
                continue
            rows.append({
                'No': values[0],
                'Jenis Alat': values[1] or '',
                'Kapasitas (Minimal)': values[2] or '',
                'Jumlah': values[3] or '',
            })
        return rows
    finally:
        wb.close()


def _format_pk_personil_summary(excel_path: str, code: str = 'PK') -> str:
    rows = _personil_rows(excel_path, code)
    if not rows:
        return 'Belum diisi; mengikuti kebutuhan dan dokumen teknis paket.'
    return '\n'.join(
        f"{row.get('No') or index}. {row.get('Jabatan', '')}; "
        f"Sertifikat: {row.get('Sertifikat', '')}; "
        f"Pengalaman: {row.get('Pengalaman Kerja', '')}"
        for index, row in enumerate(rows, 1)
    )


def _format_pk_equipment_summary(excel_path: str) -> str:
    rows = _equipment_rows(excel_path)
    if not rows:
        return 'Belum diisi; mengikuti kebutuhan dan dokumen teknis paket.'
    return '\n'.join(
        f"{row.get('No') or index}. {row.get('Jenis Alat', '')}; "
        f"Kapasitas minimal: {row.get('Kapasitas (Minimal)', '')}; "
        f"Jumlah: {row.get('Jumlah', '')}"
        for index, row in enumerate(rows, 1)
    )


def prepare_template_output(src: str, dst: str, excel_path: str, code: str) -> tuple[int, int]:
    """Copy a template and apply the Excel-driven personil table hook."""
    if os.path.basename(src).lower() == '9. list_personil.docx':
        personil_rows = merge_personil_template(src, excel_path, dst, code)
    else:
        shutil.copy2(src, dst)
        personil_rows = 0
    # KAK uses text markers for the compact resource summary. The authoritative
    # personil table remains the dedicated 9. List_Personil.docx template.
    return personil_rows, 0


# ═══════════════════════════════════════════════════════════════════════════════
# WORD REPLACEMENT
# ═══════════════════════════════════════════════════════════════════════════════
def replace_in_paragraph(para, replacements):
    full_text = ''.join(run.text for run in para.runs)
    changed   = False
    for old, new in replacements.items():
        if old in full_text:
            full_text = full_text.replace(old, new)
            changed   = True
    if changed and para.runs:
        para.runs[0].text = full_text
        for run in para.runs[1:]:
            run.text = ''
    return changed


def _process_table(table, replacements):
    """Rekursif — handle nested tables."""
    count = 0
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if replace_in_paragraph(para, replacements):
                    count += 1
            for nested in cell.tables:
                count += _process_table(nested, replacements)
    return count


def _process_textbox_markers(path, replacements):
    """Replace marker pada w:txbxContent yang tidak diekspos python-docx."""
    from lxml import etree

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path, "r") as source:
        document = etree.fromstring(source.read("word/document.xml"))
        count = 0
        for paragraph in document.xpath(".//w:txbxContent//w:p", namespaces=ns):
            nodes = paragraph.xpath(".//w:t", namespaces=ns)
            if not nodes:
                continue
            full_text = "".join(node.text or "" for node in nodes)
            new_text = full_text
            for old, new in replacements.items():
                new_text = new_text.replace(old, new)
            if new_text == full_text:
                continue
            nodes[0].text = new_text
            for node in nodes[1:]:
                node.text = ""
            count += 1
        if not count:
            return 0
        temp_path = f"{path}.textbox.tmp"
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
            payload = etree.tostring(document, xml_declaration=True, encoding="UTF-8", standalone=True)
            for item in source.infolist():
                target.writestr(item, payload if item.filename == "word/document.xml" else source.read(item.filename))
    os.replace(temp_path, path)
    return count


def _process_header_footer_markers(path, replacements):
    """Replace placeholders in Word headers/footers, which python-docx skips."""
    from lxml import etree

    changed_parts = {}
    count = 0
    with zipfile.ZipFile(path, 'r') as source:
        for item in source.infolist():
            if not (
                item.filename.startswith('word/header')
                or item.filename.startswith('word/footer')
            ) or not item.filename.endswith('.xml'):
                continue
            root = etree.fromstring(source.read(item.filename))
            part_count = 0
            for paragraph in root.xpath('.//w:p', namespaces={'w': W_NS}):
                nodes = paragraph.xpath('.//w:t', namespaces={'w': W_NS})
                if not nodes:
                    continue
                full_text = ''.join(node.text or '' for node in nodes)
                new_text = full_text
                for old, new in replacements.items():
                    new_text = new_text.replace(old, new)
                if new_text == full_text:
                    continue
                nodes[0].text = new_text
                for node in nodes[1:]:
                    node.text = ''
                part_count += 1
            if part_count:
                changed_parts[item.filename] = etree.tostring(
                    root, xml_declaration=True, encoding='UTF-8', standalone=True
                )
                count += part_count
        if not changed_parts:
            return 0
        temp_path = f'{path}.header-footer.tmp'
        with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as target:
            for item in source.infolist():
                target.writestr(
                    item,
                    changed_parts.get(item.filename, source.read(item.filename)),
                )
    os.replace(temp_path, path)
    return count


def process_docx(path, replacements):
    from docx import Document
    doc   = Document(path)
    count = 0
    for para in doc.paragraphs:
        if replace_in_paragraph(para, replacements):
            count += 1
    for table in doc.tables:
        count += _process_table(table, replacements)
    doc.save(path)
    count += _process_textbox_markers(path, replacements)
    count += _process_header_footer_markers(path, replacements)
    return count


# ═══════════════════════════════════════════════════════════════════════════════
# F5: EXPORT PDF via Word COM
# ═══════════════════════════════════════════════════════════════════════════════
def _com_path(p: str) -> str:
    r"""Strip \\?\ prefix — Word/Excel COM tidak mengenali long path prefix."""
    return p[4:] if p.startswith('\\\\?\\') else p


def _kill_lingering_word():
    """Kill instance WINWORD.EXE yang tidak punya dokumen terbuka (sisa COM)."""
    try:
        import subprocess
        result = subprocess.run(
            ['tasklist', '/FI', 'IMAGENAME eq WINWORD.EXE', '/FO', 'CSV', '/NH'],
            capture_output=True, text=True
        )
        for line in result.stdout.strip().splitlines():
            parts = [p.strip('"') for p in line.split('","')]
            if len(parts) >= 2:
                pid = parts[1]
                subprocess.run(['taskkill', '/PID', pid, '/F'],
                               capture_output=True)
    except Exception:
        pass


def export_docx_to_pdf(docx_path: str, pdf_path: str) -> None:
    import win32com.client, pythoncom, tempfile, shutil as _sh
    pythoncom.CoInitialize()
    word = None
    # Word COM tidak bisa buka path > 260 char — pakai folder temp pendek sebagai relay
    src_abs = _com_path(os.path.abspath(docx_path))
    dst_abs = _com_path(os.path.abspath(pdf_path))
    use_tmp = len(src_abs) > 240
    if use_tmp:
        tmp_dir  = tempfile.mkdtemp()
        tmp_docx = os.path.join(tmp_dir, 'doc.docx')
        tmp_pdf  = os.path.join(tmp_dir, 'doc.pdf')
        # pakai path dengan \\?\ prefix agar shutil bisa copy path panjang
        lp_src = docx_path if docx_path.startswith('\\\\?\\') else '\\\\?\\' + src_abs
        _sh.copy2(lp_src, tmp_docx)
        open_path = tmp_docx
        out_path  = tmp_pdf
    else:
        open_path = src_abs
        out_path  = dst_abs
        tmp_dir   = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(
            open_path,
            ReadOnly=False,
            AddToRecentFiles=False,
        )
        try:
            word.ActiveWindow.DisplayGridLines = False
        except Exception:
            pass
        try:
            doc.ShowGrammaticalErrors = False
            doc.ShowSpellingErrors    = False
        except Exception:
            pass
        doc.ExportAsFixedFormat(
            OutputFileName=out_path,
            ExportFormat=17,
            OpenAfterExport=False,
            OptimizeFor=0,
            Range=0,
            IncludeDocProps=True,
            KeepIRM=True,
            CreateBookmarks=0,
            DocStructureTags=True,
            BitmapMissingFonts=True,
            UseISO19005_1=False,
        )
        doc.Close(SaveChanges=False)
        # Kalau pakai tmp, copy PDF hasil ke tujuan sebenarnya (pakai \\?\ agar path panjang OK)
        if use_tmp and os.path.exists(tmp_pdf):
            lp_dst = pdf_path if pdf_path.startswith('\\\\?\\') else '\\\\?\\' + dst_abs
            _sh.copy2(tmp_pdf, lp_dst)
    finally:
        if word:
            try:
                word.Quit(SaveChanges=False)
            except Exception:
                pass
        pythoncom.CoUninitialize()
        if use_tmp and tmp_dir and os.path.isdir(tmp_dir):
            try: _sh.rmtree(tmp_dir)
            except: pass


def _export_one(word, docx_path: str, pdf_path: str) -> None:
    """Ekspor satu docx→pdf menggunakan instance Word yang sudah buka. Relay tmp jika path >240."""
    import tempfile, shutil as _sh
    src_abs = _com_path(os.path.abspath(docx_path))
    dst_abs = _com_path(os.path.abspath(pdf_path))
    use_tmp = len(src_abs) > 240
    if use_tmp:
        tmp_dir  = tempfile.mkdtemp()
        tmp_docx = os.path.join(tmp_dir, 'doc.docx')
        tmp_pdf  = os.path.join(tmp_dir, 'doc.pdf')
        lp_src   = docx_path if docx_path.startswith('\\\\?\\') else '\\\\?\\' + src_abs
        _sh.copy2(lp_src, tmp_docx)
        open_path, out_path = tmp_docx, tmp_pdf
    else:
        open_path, out_path, tmp_dir = src_abs, dst_abs, None
    try:
        doc = word.Documents.Open(open_path, ReadOnly=False, AddToRecentFiles=False)
        try: doc.ShowGrammaticalErrors = False
        except: pass
        try: doc.ShowSpellingErrors = False
        except: pass
        doc.ExportAsFixedFormat(
            OutputFileName=out_path, ExportFormat=17, OpenAfterExport=False,
            OptimizeFor=0, Range=0, IncludeDocProps=True, KeepIRM=True,
            CreateBookmarks=0, DocStructureTags=True, BitmapMissingFonts=True,
            UseISO19005_1=False,
        )
        doc.Close(SaveChanges=False)
        if use_tmp and os.path.exists(tmp_pdf):
            lp_dst = pdf_path if pdf_path.startswith('\\\\?\\') else '\\\\?\\' + dst_abs
            _sh.copy2(tmp_pdf, lp_dst)
    finally:
        if use_tmp and tmp_dir and os.path.isdir(tmp_dir):
            try: __import__('shutil').rmtree(tmp_dir)
            except: pass


def print_semua_pdf(out_dir: str) -> list:
    """Ekspor semua .docx di out_dir ke PDF — satu instance Word untuk semua file."""
    import win32com.client, pythoncom
    errors = []
    files  = [f for f in sorted(os.listdir(out_dir))
              if f.lower().endswith('.docx') and not f.startswith('~$')]
    print(f"Print PDF: {len(files)} file di {os.path.basename(out_dir)}/")
    if not files:
        return errors
    pythoncom.CoInitialize()
    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        for f in files:
            docx = os.path.join(out_dir, f)
            pdf  = docx[:-5] + '.pdf'
            try:
                _export_one(word, docx, pdf)
                print(f"  PDF OK: {f[:-5]}.pdf")
            except Exception as e:
                errors.append(f"{f}: {e}")
                print(f"  PDF GAGAL: {f}: {e}")
    finally:
        if word:
            try: word.Quit(SaveChanges=False)
            except: pass
        pythoncom.CoUninitialize()
    return errors


# ═══════════════════════════════════════════════════════════════════════════════
# F2: VALIDASI & HIGHLIGHT
# ═══════════════════════════════════════════════════════════════════════════════
def validasi_wajib(excel_data: dict) -> list:
    return [k for k in WAJIB
            if not excel_data.get(k) or str(excel_data[k]).strip() in ('', 'None')]


def highlight_wajib(xlsm_path: str, kosong_list: list) -> None:
    """Highlight cell merah/hapus highlight via Excel COM (tidak rusak xlsm)."""
    import win32com.client, pythoncom
    pythoncom.CoInitialize()
    xl = None
    try:
        xl = win32com.client.DispatchEx("Excel.Application")
        xl.Visible = False
        xl.DisplayAlerts = False
        wb = xl.Workbooks.Open(os.path.abspath(xlsm_path), ReadOnly=False, UpdateLinks=False)
        ws = wb.Sheets('Master Data')
        kosong_set = set(kosong_list)
        for row_idx in range(3, ws.UsedRange.Rows.Count + 3):
            lbl = ws.Cells(row_idx, 1).Value
            if lbl is None:
                continue
            if str(lbl) in WAJIB:
                cell = ws.Cells(row_idx, 2)
                if str(lbl) in kosong_set:
                    cell.Interior.Color = 0xCCCCFF  # merah muda (BGR: 0xCCCCFF = RGB FFCCCC)
                else:
                    cell.Interior.ColorIndex = -4142  # xlColorIndexNone
        wb.Save()
        wb.Close(SaveChanges=False)
    except Exception as e:
        print(f"  [WARN] Highlight gagal: {e}")
    finally:
        if xl:
            try: xl.Quit()
            except: pass
        pythoncom.CoUninitialize()


# ═══════════════════════════════════════════════════════════════════════════════
# F6: LOG KE EXCEL (baris 82+)
# ═══════════════════════════════════════════════════════════════════════════════
def tulis_log_excel(xlsm_path: str, log_data: dict) -> None:
    """Tulis log generate ke ppk_log.xlsx (bukan ke xlsm agar tidak korupsi VBA)."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill
        if os.path.exists(LOG_XLSX):
            wb = openpyxl.load_workbook(LOG_XLSX)
        else:
            wb = openpyxl.Workbook()
            # hapus sheet default
            if 'Sheet' in wb.sheetnames:
                del wb['Sheet']

        if 'Log Generate' not in wb.sheetnames:
            ws = wb.create_sheet('Log Generate')
            headers = ['Waktu', 'Folder Output', 'Total Replace', 'Field Kosong', 'Status']
            for i, h in enumerate(headers, 1):
                c = ws.cell(1, i, h)
                c.font = Font(bold=True, color='FFFFFF')
                c.fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
        else:
            ws = wb['Log Generate']

        next_row = ws.max_row + 1
        kosong = log_data.get('field_kosong', [])
        ws.cell(next_row, 1, log_data.get('waktu', ''))
        ws.cell(next_row, 2, log_data.get('folder', ''))
        ws.cell(next_row, 3, log_data.get('total_replace', 0))
        ws.cell(next_row, 4, ', '.join(kosong) if isinstance(kosong, list) else str(kosong))
        status = log_data.get('status', '')
        sc = ws.cell(next_row, 5, status)
        if 'SUKSES' in status:
            sc.font = Font(bold=True, color='006400')
        elif any(x in status for x in ('GAGAL', 'ERROR')):
            sc.font = Font(bold=True, color='CC0000')
        wb.save(LOG_XLSX)
    except Exception as e:
        print(f"  [WARN] Tulis log gagal: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
# F8: CATAT RIWAYAT
# ═══════════════════════════════════════════════════════════════════════════════
def catat_riwayat(xlsm_path: str, excel_data: dict, out_dir: str, status: str) -> None:
    """Catat riwayat ke ppk_log.xlsx sheet Riwayat."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill
        if os.path.exists(LOG_XLSX):
            wb = openpyxl.load_workbook(LOG_XLSX)
        else:
            wb = openpyxl.Workbook()
            if 'Sheet' in wb.sheetnames:
                del wb['Sheet']

        if 'Riwayat' not in wb.sheetnames:
            ws = wb.create_sheet('Riwayat')
            headers = ['No', 'Nama Paket (Lengkap)', 'Kode RUP', 'Pagu Anggaran', 'Tanggal Generate', 'Folder Output', 'Status']
            col_widths = [5, 45, 15, 18, 20, 40, 30]
            for i, (h, w) in enumerate(zip(headers, col_widths), 1):
                c = ws.cell(1, i, h)
                c.font = Font(bold=True, color='FFFFFF')
                c.fill = PatternFill(start_color='7B2D8B', end_color='7B2D8B', fill_type='solid')
                from openpyxl.utils import get_column_letter
                ws.column_dimensions[get_column_letter(i)].width = w
        else:
            ws = wb['Riwayat']

        next_row = ws.max_row + 1
        def _int(v):
            try: return int(str(v).replace('.', '').replace(',', '').strip())
            except: return 0
        ws.cell(next_row, 1, next_row - 1)
        ws.cell(next_row, 2, excel_data.get('Nama Paket (Lengkap)', ''))
        ws.cell(next_row, 3, str(excel_data.get('Kode RUP', '')))
        ws.cell(next_row, 4, _int(excel_data.get('Pagu Anggaran (Angka)', 0)))
        ws.cell(next_row, 5, datetime.now().strftime('%Y-%m-%d %H:%M'))
        ws.cell(next_row, 6, os.path.basename(out_dir))
        ws.cell(next_row, 7, status)
        wb.save(LOG_XLSX)
    except Exception as e:
        print(f"  [WARN] Catat riwayat gagal: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
# F9: SUPABASE SYNC
# ═══════════════════════════════════════════════════════════════════════════════
def _sb():
    from supabase import create_client
    from dotenv import load_dotenv
    load_dotenv(SB_ENV)
    return create_client(os.environ['SUPABASE_URL'], os.environ['SUPABASE_KEY'])


def sync_supabase(excel_data: dict) -> dict:
    kode = str(excel_data.get('Kode RUP', '')).strip()
    if not kode or kode == 'None':
        return {'ok': False, 'error': 'Kode RUP kosong'}
    def _int(v):
        try:
            return int(str(v).replace('.', '').replace(',', '').strip())
        except Exception:
            return 0
    data = {
        'kode_paket':    kode,
        'nama_paket':    excel_data.get('Nama Paket (Lengkap)', ''),
        'nilai_hps':     _int(excel_data.get('Nilai HPS (Angka)', 0)),
        'pagu':          _int(excel_data.get('Pagu Anggaran (Angka)', 0)),
        'jenis_pl':      'JKK',
        'folder_dibuat': _re.sub(r'[\\/:*?"<>|]', '', excel_data.get('Nama Paket (Lengkap)', '')),
    }
    try:
        _sb().table('draft_paket_pl').upsert(data, on_conflict='kode_paket').execute()
        return {'ok': True}
    except Exception as e:
        return {'ok': False, 'error': str(e)}


# ═══════════════════════════════════════════════════════════════════════════════
# BACA EXCEL
# ═══════════════════════════════════════════════════════════════════════════════
def _format_lingkup_pekerjaan(items) -> str:
    """Format maksimal 10 item lingkup secara inline dengan konjungsi natural."""
    if isinstance(items, str):
        raw_items = items.splitlines()
    else:
        try:
            raw_items = list(items or [])
        except TypeError:
            raw_items = [items]

    cleaned = []
    for item in raw_items:
        value = str(item or '').strip()
        if not value:
            continue
        # Hindari nomor ganda bila data sudah pernah diformat sebelumnya.
        value = _re.sub(r'^\s*\d+\s*[.)]\s*', '', value).strip()
        if value:
            cleaned.append(value)

    numbered = [f'{index}. {value}' for index, value in enumerate(cleaned, 1)]
    if len(numbered) <= 1:
        return numbered[0] if numbered else ''
    if len(numbered) == 2:
        return f'{numbered[0]}, dan {numbered[1]}'
    return ', '.join(numbered[:-1]) + f', dan {numbered[-1]}'


def baca_excel(xlsm_path: str) -> dict:
    import openpyxl
    wb = openpyxl.load_workbook(xlsm_path, data_only=True)
    ws = wb['Master Data']
    data = {}
    for row in ws.iter_rows(min_row=3, values_only=True):
        label = row[0]
        if label and not str(label).startswith('---') and not str(label).startswith('==='):
            data[str(label)] = str(row[1]) if row[1] is not None else ''
    scope = []
    for row_number in range(14, 24):
        value = ws.cell(row_number, 5).value
        if value is not None and str(value).strip():
            scope.append(str(value).strip())
    data["Lingkup Pekerjaan"] = _format_lingkup_pekerjaan(scope)
    return data


# ═══════════════════════════════════════════════════════════════════════════════
# F1: AUTO-TERBILANG (isi field Terbilang yang kosong dari field Angka)
# ═══════════════════════════════════════════════════════════════════════════════
def auto_terbilang(excel_data: dict) -> dict:
    pairs = [
        ('Pagu Anggaran (Angka)',   'Pagu Anggaran (Terbilang)',   'rupiah'),
        ('Nilai HPS (Angka)',       'Nilai HPS (Terbilang)',        'rupiah'),
        ('Nilai Penawaran (Angka)', 'Nilai Penawaran (Terbilang)',  'rupiah'),
        ('Nilai Kontrak (Angka)',   'Nilai Kontrak (Terbilang)',    'rupiah'),
    ]
    for k_angka, k_terbilang, tipe in pairs:
        angka = excel_data.get(k_angka, '').strip()
        if angka and angka != 'None':
            if not excel_data.get(k_terbilang) or excel_data[k_terbilang].strip() in ('', 'None'):
                excel_data[k_terbilang] = terbilang_rupiah(angka)
                print(f"  Auto-terbilang: [{k_terbilang}] = {excel_data[k_terbilang]}")

    hari = excel_data.get('Jangka Waktu (Hari)', '').strip()
    if hari and hari != 'None':
        if not excel_data.get('Jangka Waktu (Terbilang)') or excel_data['Jangka Waktu (Terbilang)'].strip() in ('', 'None'):
            excel_data['Jangka Waktu (Terbilang)'] = terbilang_hari(hari)
            print(f"  Auto-terbilang: [Jangka Waktu (Terbilang)] = {excel_data['Jangka Waktu (Terbilang)']}")
    return _auto_tanggal_selesai(excel_data)


# ═══════════════════════════════════════════════════════════════════════════════
# BUILD REPLACEMENTS
# ═══════════════════════════════════════════════════════════════════════════════
def build_replacements(
    excel_data: dict,
    document_name: str = "",
    excel_path: str | None = None,
) -> dict:
    source_excel_path = excel_path or EXCEL_PATH
    stage = normalize_document_stage(excel_data.get("Tahap Dokumen"))
    kind = contract_doc_kind(document_name)
    detail = detail_enabled(stage, kind)
    wakil_status = str(excel_data.get("Ada Wakil Sah?", "") or "").strip().upper()
    repl = {}
    for excel_key, placeholder in FIELD_MAP.items():
        # Marker ini adalah heading blok donor standar. Pertahankan heading
        # meskipun tidak ada wakil, agar blok opsional tidak tampak terhapus.
        if excel_key == "Ada Wakil Sah?":
            repl[placeholder] = "WAKIL SAH PEJABAT PENANDATANGAN KONTRAK"
            continue
        if not detail and excel_key in CONTRACT_FIELD_KEYS:
            repl[placeholder] = ""
            continue
        val = excel_data.get(excel_key, '')
        if excel_key in {
            "Nama Wakil Sah", "Jabatan Wakil Sah", "NIP Wakil Sah",
            "Nomor SK Wakil Sah", "Tanggal SK Wakil Sah",
        }:
            if wakil_status != "ADA":
                val = ""
        if excel_key == "Nama Wakil Sah" and wakil_status == "ADA":
            nama_wakil = str(excel_data.get("Nama Wakil Sah", "") or "").strip()
            jabatan_wakil = str(excel_data.get("Jabatan Wakil Sah", "") or "").strip()
            val = " — ".join(part for part in (nama_wakil, jabatan_wakil) if part)
        if excel_key == "Uang Muka (%)" and stage == "BERKONTRAK" and not str(val or "").strip():
            val = f"{_minimum_uang_muka(excel_data):g}%"
        elif excel_key in {"Uang Muka (%)", "Retensi (%)"} and str(val or "").strip():
            val = _format_percentage(val)
        if val and val != 'None':
            repl[placeholder] = val
        else:
            repl[placeholder] = ""
    # Format angka dengan titik ribuan («PAGU_ANGKA_FORMAT», dll.)
    for key, (ph_raw, ph_fmt) in {
        'Pagu Anggaran (Angka)':   ('«PAGU_ANGKA»',            '«PAGU_ANGKA_FORMAT»'),
        'Nilai HPS (Angka)':       ('«HPS_ANGKA»',              '«HPS_ANGKA_FORMAT»'),
        'Nilai Penawaran (Angka)': ('«NILAI_PENAWARAN_ANGKA»',  '«NILAI_PENAWARAN_ANGKA_FORMAT»'),
        'Nilai Kontrak (Angka)':   ('«NILAI_KONTRAK_ANGKA»',    '«NILAI_KONTRAK_ANGKA_FORMAT»'),
    }.items():
        if not detail and key in CONTRACT_FIELD_KEYS:
            repl[ph_raw] = ''
            repl[ph_fmt] = ''
            continue
        val = excel_data.get(key, '')
        if val and val != 'None':
            try:
                fmt = f"{int(str(val).replace('.','').replace(',','').strip()):,}".replace(',', '.')
            except Exception:
                fmt = val
            repl[ph_raw] = val
            repl[ph_fmt] = fmt
        else:
            repl[ph_raw] = ''
            repl[ph_fmt] = ''
    # SPPBJ upload awal tetap informatif, tetapi nominal belum ditetapkan.
    # Saat BERKONTRAK, kedua varian donor memakai nilai kontrak.
    if kind == "sppbj":
        sppbj_placeholders = (
            ("\u00abNILAI_PENAWARAN_ANGKA_FORMAT\u00bb", "\u00abNILAI_PENAWARAN_TERBILANG\u00bb"),
            ("\u00abNILAI_KONTRAK_ANGKA_FORMAT\u00bb", "\u00abNILAI_KONTRAK_TERBILANG\u00bb"),
        )
        if stage == "UPLOAD AWAL":
            for ph_number, ph_words in sppbj_placeholders:
                repl[ph_number] = ""
                repl[ph_words] = "____"
        else:
            contract_number = repl.get("\u00abNILAI_KONTRAK_ANGKA_FORMAT\u00bb", "")
            contract_words = str(excel_data.get("Nilai Kontrak (Terbilang)", "") or "").strip()
            for ph_number, ph_words in sppbj_placeholders:
                repl[ph_number] = contract_number
                repl[ph_words] = contract_words

    if not detail:
        # Dokumen upload awal tetap menunjukkan lokasi input kontrak dengan
        # underscore; field substantif tertentu tetap ditampilkan apa adanya.
        upload_blank_fields = CONTRACT_FIELD_KEYS - {
            "Masa Pemeliharaan (Hari)", "Sistem Pembayaran",
            "Ada Wakil Sah?",
            "Nilai Penawaran (Angka)", "Nilai Penawaran (Terbilang)",
            "Nilai Kontrak (Angka)", "Nilai Kontrak (Terbilang)",
        }
        for excel_key in upload_blank_fields:
            placeholder = FIELD_MAP.get(excel_key)
            if placeholder:
                repl[placeholder] = "____"
        masa_pemeliharaan = str(
            excel_data.get("Masa Pemeliharaan (Hari)", "") or ""
        ).strip()
        if masa_pemeliharaan:
            repl[FIELD_MAP["Masa Pemeliharaan (Hari)"]] = masa_pemeliharaan

    # Placeholder blok dinamis: rekening bank dan harga kontrak hanya pada
    # BERKONTRAK. Sistem pembayaran tetap ditampilkan sejak UPLOAD AWAL.
    contract = stage == "BERKONTRAK"
    repl["\u00abHARGA_KONTRAK\u00bb"] = ""
    if contract and repl.get("\u00abNILAI_KONTRAK_ANGKA_FORMAT\u00bb"):
        repl["\u00abHARGA_KONTRAK\u00bb"] = (
            "Harga Kontrak termasuk Pajak Pertambahan Nilai (PPN) adalah sebesar "
            f"Rp{repl['\u00abNILAI_KONTRAK_ANGKA_FORMAT\u00bb']} "
            f"({repl.get('«NILAI_KONTRAK_TERBILANG»', '')} rupiah) yang diperoleh "
            "berdasarkan total harga penawaran terkoreksi aritmatik sebagaimana "
            "tercantum dalam Daftar Kuantitas dan Harga hasil negosiasi dan koreksi aritmatik."
        )
    elif kind == "spk":
        repl["\u00abHARGA_KONTRAK\u00bb"] = "(Rp. _____)"

    repl["\u00abPEMBAYARAN_BANK\u00bb"] = ""
    if contract:
        repl["\u00abPEMBAYARAN_BANK\u00bb"] = (
            "Pembayaran untuk kontrak ini dilakukan ke "
            f"{excel_data.get('Nama Bank', '')} rekening nomor : "
            f"{excel_data.get('Nomor Rekening', '')} atas nama Penyedia : "
            f"{excel_data.get('Atas Nama Rekening', '')}"
        )
    elif kind == "spk":
        repl["\u00abPEMBAYARAN_BANK\u00bb"] = (
            "Pembayaran untuk kontrak ini dilakukan ke Bank ____ rekening nomor : "
            "____ atas nama Penyedia : ____"
        )

    repl["\u00abKETENTUAN_UANG_MUKA\u00bb"] = (
        "Untuk nilai HPS s/d Rp. 200.000.000,00 (Dua Ratus Juta Rupiah) "
        "diberikan uang muka paling sedikit sebesar 50% dari Nilai Kontrak "
        "dan untuk nilai HPS di atas Rp. 200.000.000,00 (Dua Ratus Juta Rupiah) "
        "diberikan uang muka paling sedikit sebesar 30% dari Nilai Kontrak."
    )
    repl["\u00abUANG_MUKA_KONTRAK\u00bb"] = ""
    if contract and repl.get("\u00abUANG_MUKA_PERSEN\u00bb"):
        repl["\u00abUANG_MUKA_KONTRAK\u00bb"] = (
            "Kontrak ini diberikan uang muka sebesar "
            f"{repl['\u00abUANG_MUKA_PERSEN\u00bb']} dari Nilai Kontrak"
        )
    elif kind == "spk":
        repl["\u00abUANG_MUKA_KONTRAK\u00bb"] = (
            "Kontrak ini diberikan uang muka sebesar ____% dari Nilai Kontrak"
        )

    payment = str(excel_data.get("Sistem Pembayaran", "") or "").strip()
    retensi = repl.get("\u00abRETENSI_PERSEN\u00bb", "") if contract else ""
    repl["\u00abPEMBAYARAN_PRESTASI\u00bb"] = ""
    if payment:
        repl["\u00abPEMBAYARAN_PRESTASI\u00bb"] = (
            "Pembayaran prestasi pekerjaan dilakukan dengan cara : "
            f"{payment}"
            + (f"; Retensi: {retensi}" if retensi else "")
        )
    repl["\u00abRINCIAN_TERMIN_BLOK\u00bb"] = ""
    if contract and str(excel_data.get("Rincian Termin", "") or "").strip():
        repl["\u00abRINCIAN_TERMIN_BLOK\u00bb"] = (
            "Rincian pembayaran/termin: "
            f"{excel_data.get('Rincian Termin', '')}"
        )
    elif kind == "spk" and payment:
        repl["\u00abRINCIAN_TERMIN_BLOK\u00bb"] = "Rincian pembayaran/termin: ____"

    repl["\u00abSUMBER_DANA_DETAIL\u00bb"] = _source_funds_detail(excel_data)
    # Blok tanda tangan penyedia tetap informatif pada UPLOAD AWAL, tetapi
    # memakai identitas final saat BERKONTRAK. Marker khusus mencegah
    # perubahan label identitas penyedia/direktur di badan dokumen.
    repl["\u00abNAMA_PENYEDIA_TTD\u00bb"] = (
        "[Nama Penyedia]"
        if not detail
        else str(excel_data.get("Nama Direktur", "") or "").strip()
    )
    repl["\u00abDIREKTUR_TTD\u00bb"] = (
        "[Direktur]"
        if not detail
        else str(excel_data.get("Jabatan Direktur", "") or "").strip()
    )
    repl["\u00abLINGKUP_PEKERJAAN\u00bb"] = _format_lingkup_pekerjaan(
        excel_data.get("Lingkup Pekerjaan", "")
    )
    if str(excel_data.get("Kode Jenis Paket", "") or "").strip().upper() == "PK":
        repl["\u00abTABEL_PERSONEL_PK\u00bb"] = _format_pk_personil_summary(source_excel_path, "PK")
        repl["\u00abTABEL_PERALATAN_PK\u00bb"] = _format_pk_equipment_summary(source_excel_path)
    for key in ("Nomor SK PPK", "Tanggal SK PPK", "Uraian SK PPK"):
        repl[FIELD_MAP[key]] = str(excel_data.get(key, "") or "")
    return repl


# ═══════════════════════════════════════════════════════════════════════════════
# F7: NOMOR BARU (auto-increment dari sheet Nomor Surat)
# ═══════════════════════════════════════════════════════════════════════════════
_JENIS_TO_FIELD = {
    'Nota Dinas':     'Nomor Nota Dinas',
    'SPPBJ':          'Nomor SPPBJ',
    'SPK':            'Nomor SPK',
    'SPMK':           'Nomor SPMK',
    'Surat Undangan': 'Nomor Surat Undangan PL',
    'BA Hasil PL':    'Nomor BA Hasil PL',
}


def nomor_baru(xlsm_path: str, jenis: str) -> dict:
    """Auto-increment nomor surat via Excel COM (aman untuk xlsm)."""
    import win32com.client, pythoncom
    if jenis not in _JENIS_TO_FIELD:
        return {'ok': False, 'error': f"Jenis tidak dikenal: '{jenis}'. Pilihan: {list(_JENIS_TO_FIELD)}"}
    pythoncom.CoInitialize()
    xl = None
    try:
        xl = win32com.client.DispatchEx("Excel.Application")
        xl.Visible = False
        xl.DisplayAlerts = False
        wb = xl.Workbooks.Open(os.path.abspath(xlsm_path), ReadOnly=False, UpdateLinks=False)

        # Cek sheet Nomor Surat ada
        sheet_names = [wb.Sheets(i+1).Name for i in range(wb.Sheets.Count)]
        if 'Nomor Surat' not in sheet_names:
            wb.Close(SaveChanges=False)
            return {'ok': False, 'error': "Sheet 'Nomor Surat' belum ada. Jalankan inject_vba_ppk.py dulu."}

        ws_nomor = wb.Sheets('Nomor Surat')
        ws_master = wb.Sheets('Master Data')

        # Cari baris jenis di sheet Nomor Surat (mulai baris 2)
        target_row = None
        for r in range(2, 20):
            val = ws_nomor.Cells(r, 1).Value
            if val == jenis:
                target_row = r
                break
        if target_row is None:
            wb.Close(SaveChanges=False)
            return {'ok': False, 'error': f"Jenis '{jenis}' tidak ada di sheet Nomor Surat."}

        prefix = str(ws_nomor.Cells(target_row, 2).Value or '')
        tahun = str(ws_nomor.Cells(target_row, 3).Value or datetime.now().year)
        nomor_lama = int(ws_nomor.Cells(target_row, 4).Value or 0)
        nomor_baru_n = nomor_lama + 1
        nomor_str = f"{prefix}{nomor_baru_n:02d}/{tahun}"

        ws_nomor.Cells(target_row, 4).Value = nomor_baru_n
        ws_nomor.Cells(target_row, 5).Value = nomor_str

        # Update field di Master Data
        field_master = _JENIS_TO_FIELD[jenis]
        for r in range(3, ws_master.UsedRange.Rows.Count + 3):
            lbl = ws_master.Cells(r, 1).Value
            if lbl == field_master:
                ws_master.Cells(r, 2).Value = nomor_str
                break

        wb.Save()
        wb.Close(SaveChanges=False)
        print(f"  Nomor baru [{jenis}]: {nomor_str}")
        return {'ok': True, 'nomor': nomor_str}
    except Exception as e:
        return {'ok': False, 'error': str(e)}
    finally:
        if xl:
            try: xl.Quit()
            except: pass
        pythoncom.CoUninitialize()


# ═══════════════════════════════════════════════════════════════════════════════
# F10: PAKET BARU (reset field input)
# ═══════════════════════════════════════════════════════════════════════════════
def paket_baru(xlsm_path: str) -> None:
    """Reset field input untuk paket baru via Excel COM."""
    import win32com.client, pythoncom
    pythoncom.CoInitialize()
    xl = None
    try:
        xl = win32com.client.DispatchEx("Excel.Application")
        xl.Visible = False
        xl.DisplayAlerts = False
        wb = xl.Workbooks.Open(os.path.abspath(xlsm_path), ReadOnly=False, UpdateLinks=False)
        ws = wb.Sheets('Master Data')
        reset_set = set(RESET_FIELDS)
        count = 0
        for r in range(3, ws.UsedRange.Rows.Count + 3):
            lbl = ws.Cells(r, 1).Value
            if lbl and str(lbl) in reset_set:
                ws.Cells(r, 2).Value = None
                count += 1
            if str(lbl or '').strip() == 'Tahap Dokumen':
                ws.Cells(r, 2).Value = 'UPLOAD AWAL'
        for row_number in range(14, 24):
            ws.Cells(row_number, 5).Value = None
        wb.Save()
        wb.Close(SaveChanges=False)
        print(f"  Reset selesai: {count} field direset. Seksi B (SKPD) & C (PPK) dipertahankan.")
    except Exception as e:
        print(f"ERROR paket_baru: {e}")
        import traceback; traceback.print_exc()
    finally:
        if xl:
            try: xl.Quit()
            except: pass
        pythoncom.CoUninitialize()


# ═══════════════════════════════════════════════════════════════════════════════
# F3: COMMIT PAKET ke sheet Daftar Paket
# ═══════════════════════════════════════════════════════════════════════════════
def commit_paket(xlsm_path: str, excel_data: dict) -> None:
    """Salin data paket ke sheet Daftar Paket via Excel COM."""
    import win32com.client, pythoncom
    pythoncom.CoInitialize()
    xl = None
    try:
        xl = win32com.client.DispatchEx("Excel.Application")
        xl.Visible = False
        xl.DisplayAlerts = False
        wb = xl.Workbooks.Open(os.path.abspath(xlsm_path), ReadOnly=False, UpdateLinks=False)
        sheet_names = [wb.Sheets(i+1).Name for i in range(wb.Sheets.Count)]
        if 'Daftar Paket' not in sheet_names:
            wb.Close(SaveChanges=False)
            print("  [WARN] Sheet 'Daftar Paket' belum ada. Jalankan inject_vba_ppk.py dulu.")
            return
        ws = wb.Sheets('Daftar Paket')
        # Cari baris kosong berikutnya (setelah baris terakhir berisi data)
        last_row = ws.UsedRange.Rows.Count + 1 if ws.UsedRange.Rows.Count > 1 else 2
        # Pastikan baris 1 sudah diisi (header), start dari baris 2
        next_row = max(2, last_row)
        def _int(v):
            try: return int(str(v).replace('.', '').replace(',', '').strip())
            except: return 0
        ws.Cells(next_row, 1).Value = next_row - 1
        ws.Cells(next_row, 2).Value = excel_data.get('Nama Paket (Lengkap)', '')
        ws.Cells(next_row, 3).Value = excel_data.get('Nama Paket (Lengkap)', '')
        ws.Cells(next_row, 4).Value = str(excel_data.get('Kode RUP', ''))
        ws.Cells(next_row, 5).Value = _int(excel_data.get('Pagu Anggaran (Angka)', 0))
        ws.Cells(next_row, 6).Value = _int(excel_data.get('Nilai HPS (Angka)', 0))
        ws.Cells(next_row, 7).Value = excel_data.get('Lokasi Pekerjaan', '')
        ws.Cells(next_row, 8).Value = excel_data.get('Tanggal KAK & HPS', '')
        ws.Cells(next_row, 9).Value = excel_data.get('Nama Penyedia', '')
        ws.Cells(next_row, 10).Value = ''
        wb.Save()
        wb.Close(SaveChanges=False)
        print(f"  Paket di-commit ke 'Daftar Paket' baris {next_row}.")
    except Exception as e:
        print(f"  [WARN] Commit paket gagal: {e}")
    finally:
        if xl:
            try: xl.Quit()
            except: pass
        pythoncom.CoUninitialize()


# ═══════════════════════════════════════════════════════════════════════════════
# F3: GENERATE MULTI-PAKET
# ═══════════════════════════════════════════════════════════════════════════════
def generate_multi(xlsm_path: str) -> None:
    """Generate dokumen untuk semua paket di sheet Daftar Paket."""
    import openpyxl, win32com.client, pythoncom

    # Baca Daftar Paket via openpyxl read-only (AMAN — tidak save)
    wb_ro = openpyxl.load_workbook(xlsm_path, data_only=True, read_only=True)
    sheet_names_ro = wb_ro.sheetnames
    if 'Daftar Paket' not in sheet_names_ro:
        wb_ro.close()
        print("ERROR: Sheet 'Daftar Paket' belum ada. Jalankan inject_vba_ppk.py dulu.")
        sys.exit(1)
    ws_dp = wb_ro['Daftar Paket']
    ws_md = wb_ro['Master Data']

    # Baca base_data dari Master Data
    base_data = {}
    for row in ws_md.iter_rows(min_row=3, values_only=True):
        if row[0] and not str(row[0]).startswith('---') and not str(row[0]).startswith('==='):
            base_data[str(row[0])] = str(row[1]) if row[1] is not None else ''

    # Baca semua baris Daftar Paket
    COL_FIELD = {
        3: 'Nama Paket (Lengkap)',
        4: 'Kode RUP',            5: 'Pagu Anggaran (Angka)',
        6: 'Nilai HPS (Angka)',   7: 'Lokasi Pekerjaan',
        8: 'Tanggal KAK & HPS',   9: 'Nama Penyedia',
    }
    paket_list = []  # list of (row_idx, excel_data)
    rows = list(ws_dp.iter_rows(min_row=2, values_only=True))
    wb_ro.close()

    for i, row_vals in enumerate(rows):
        row_idx = i + 2
        if not row_vals[3]:  # kolom 4 = Kode RUP (index 3)
            continue
        status_lama = str(row_vals[9] or '')  # kolom 10 = Status (index 9)
        if status_lama.startswith('SUKSES'):
            print(f"  SKIP baris {row_idx}: sudah di-generate.")
            continue
        excel_data = dict(base_data)
        for col_idx, field_name in COL_FIELD.items():
            val = row_vals[col_idx - 1]  # row_vals is 0-indexed
            excel_data[field_name] = str(val) if val is not None else ''
        # Fallback nomor urut dari index loop jika belum di-CommitPaket
        no_urut = excel_data.get('Nomor Urut Paket', '').strip()
        if not no_urut or no_urut == 'None':
            excel_data['Nomor Urut Paket'] = str(i + 1)
        paket_list.append((row_idx, excel_data))

    if not paket_list:
        print("Tidak ada paket baru untuk di-generate.")
        return

    print(f"Multi-generate: {len(paket_list)} paket")

    # Generate dokumen + update status via COM
    pythoncom.CoInitialize()
    xl = None
    try:
        xl = win32com.client.DispatchEx("Excel.Application")
        xl.Visible = False
        xl.DisplayAlerts = False
        wb = xl.Workbooks.Open(os.path.abspath(xlsm_path), ReadOnly=False, UpdateLinks=False)
        ws_com = wb.Sheets('Daftar Paket')

        for row_idx, excel_data in paket_list:
            excel_data = auto_terbilang(excel_data)
            _nf = excel_data.get('Nama Paket (Lengkap)', '').strip()
            safe_nama = _re.sub(r'[\\/:*?"<>|]', '', _nf).strip()
            _base_lp = '\\\\?\\' + BASE_DIR
            out_dir = os.path.join(_base_lp, safe_nama or f'OUTPUT_{datetime.now().strftime("%Y%m%d_%H%M%S")}')
            os.makedirs(out_dir, exist_ok=True)
            _kode_template, _folder_template = _template_profile(excel_data)
            word_files = template_files_for_data(BASE_DIR, excel_data)
            filled = 0
            for src_fname, out_prefix in word_files:
                src = os.path.join(BASE_DIR, src_fname)
                if not os.path.exists(src):
                    continue
                dst_fname = f"{out_prefix}.docx" if out_prefix else src_fname
                dst = os.path.join(out_dir, dst_fname)
                _rows, _resource_tables = prepare_template_output(
                    src, dst, xlsm_path, _kode_template
                )
                repl = build_replacements(excel_data, src_fname, xlsm_path)
                filled += process_docx(dst, repl)
            status_str = f'SUKSES ({filled} replace) — {datetime.now().strftime("%Y-%m-%d %H:%M")}'
            ws_com.Cells(row_idx, 10).Value = status_str
            print(f"  Baris {row_idx}: [{safe_nama}] -> {filled} replace")

        wb.Save()
        wb.Close(SaveChanges=False)
    except Exception as e:
        print(f"ERROR multi-generate: {e}")
        import traceback; traceback.print_exc()
    finally:
        if xl:
            try: xl.Quit()
            except: pass
        pythoncom.CoUninitialize()

    print("\nMulti-generate selesai.")


# ═══════════════════════════════════════════════════════════════════════════════
# F11: MUAT DARI DB
# ═══════════════════════════════════════════════════════════════════════════════
def tarik_dari_supabase(kode_rup: str) -> dict | None:
    """Tarik data paket dari Supabase berdasarkan kode_rup."""
    try:
        result = _sb().table('draft_paket_pl').select('*').eq('kode_paket', kode_rup).execute()
        return result.data[0] if result.data else None
    except Exception as e:
        print(f"ERROR Supabase: {e}")
        return None


def isi_excel_dari_supabase(xlsm_path: str, data: dict) -> None:
    """Tulis data dari Supabase ke Excel via COM. Ikut pola paket_baru()."""
    import win32com.client, pythoncom
    # Mapping kolom Supabase → label di kolom A Excel sheet Master Data
    SB_TO_EXCEL = {
        'nama_paket':       'Nama Paket (Lengkap)',
        'kode_paket':       'Kode RUP',
        'pagu':             'Pagu Anggaran (Angka)',
        'lokasi_pekerjaan': 'Lokasi Pekerjaan',
        'sumber_dana':      'Sumber Dana',
        'mak':              'Kode Rekening (MAK)',
        'tahun_anggaran':   'Tahun Anggaran',
        'nama_ppk':         'Nama PPK',
        'nip_ppk':          'NIP PPK',
    }
    pythoncom.CoInitialize()
    xl = None
    try:
        xl = win32com.client.DispatchEx("Excel.Application")
        xl.Visible = False
        xl.DisplayAlerts = False
        wb = xl.Workbooks.Open(os.path.abspath(xlsm_path), ReadOnly=False, UpdateLinks=False)
        ws = wb.Sheets('Master Data')
        filled = 0
        for r in range(3, ws.UsedRange.Rows.Count + 3):
            lbl = ws.Cells(r, 1).Value
            if lbl is None:
                continue
            for sb_col, excel_label in SB_TO_EXCEL.items():
                if str(lbl) == excel_label:
                    val = data.get(sb_col)
                    if val is not None and str(val) not in ('', 'None'):
                        ws.Cells(r, 2).Value = str(val)
                        filled += 1
        wb.Save()
        wb.Close(SaveChanges=False)
        print(f"  {filled} field diisi dari Supabase.")
    except Exception as e:
        print(f"ERROR isi_excel_dari_supabase: {e}")
        import traceback; traceback.print_exc()
    finally:
        if xl:
            try: xl.Quit()
            except: pass
        pythoncom.CoUninitialize()


# ═══════════════════════════════════════════════════════════════════════════════
# F12: SNAPSHOT (SIMPAN & MUAT)
# ═══════════════════════════════════════════════════════════════════════════════
def simpan_snapshot(xlsm_path: str, kode_rup: str) -> dict:
    import win32com.client, pythoncom, json
    pythoncom.CoInitialize()
    xl = None
    try:
        xl = win32com.client.DispatchEx("Excel.Application")
        xl.Visible = False
        xl.DisplayAlerts = False
        wb = xl.Workbooks.Open(os.path.abspath(xlsm_path), ReadOnly=False, UpdateLinks=False)

        # 1. Baca semua field dari Master Data
        ws_master = wb.Sheets('Master Data')
        fields = {}
        for r in range(3, ws_master.UsedRange.Rows.Count + 3):
            lbl = ws_master.Cells(r, 1).Value
            if lbl is not None:
                lbl_str = str(lbl).strip()
                if lbl_str and not lbl_str.startswith('---') and not lbl_str.startswith('==='):
                    val = ws_master.Cells(r, 2).Value
                    fields[lbl_str] = str(val) if val is not None else ''
        fields["Lingkup Pekerjaan"] = [
            str(ws_master.Cells(row_number, 5).Value or "")
            for row_number in range(14, 24)
        ]

        # 2. Cari baris di Daftar Paket yang kode_rup-nya cocok (kolom 4)
        ws_dp = wb.Sheets('Daftar Paket')

        # Pastikan kolom 11 ada header "Snapshot"
        hdr = ws_dp.Cells(1, 11).Value
        if hdr is None or str(hdr).strip() != 'Snapshot':
            ws_dp.Cells(1, 11).Value = 'Snapshot'
            ws_dp.Columns(11).ColumnWidth = 30

        found_row = None
        def _clean(v):
            if v is None: return ""
            s = str(v).strip()
            if s.endswith(".0"): s = s[:-2]
            return s

        target_rup = _clean(kode_rup)
        for r in range(2, ws_dp.UsedRange.Rows.Count + 2):
            val_kode = ws_dp.Cells(r, 4).Value
            if val_kode is not None and _clean(val_kode) == target_rup:
                found_row = r
                break

        if found_row is None:
            wb.Close(SaveChanges=False)
            return {'ok': False, 'error': f"Kode RUP '{kode_rup}' tidak ditemukan di sheet Daftar Paket."}

        # 3. Tulis JSON
        json_data = json.dumps(fields, ensure_ascii=False)
        ws_dp.Cells(found_row, 11).Value = json_data

        wb.Save()
        wb.Close(SaveChanges=False)
        print(f"  Snapshot disimpan ke sheet 'Daftar Paket' baris {found_row} untuk RUP {kode_rup}")
        return {'ok': True}
    except Exception as e:
        return {'ok': False, 'error': str(e)}
    finally:
        if xl:
            try: xl.Quit()
            except: pass
        pythoncom.CoUninitialize()


def muat_snapshot(xlsm_path: str, kode_rup: str) -> dict:
    import win32com.client, pythoncom, json
    pythoncom.CoInitialize()
    abs_path = os.path.abspath(xlsm_path)

    def _path_key(path):
        """Samakan path junction C: dengan target G: sebelum attach COM."""
        try:
            return os.path.normcase(os.path.realpath(os.path.abspath(str(path))))
        except (OSError, TypeError, ValueError):
            return os.path.normcase(os.path.abspath(str(path)))

    target_path_key = _path_key(abs_path)

    def _clean(v):
        if v is None: return ""
        s = str(v).strip()
        if s.endswith(".0"): s = s[:-2]
        return s

    def _do_restore(wb):
        # 1. Cari baris di Daftar Paket
        ws_dp = wb.Sheets('Daftar Paket')
        found_row = None
        target_rup = _clean(kode_rup)
        for r in range(2, ws_dp.UsedRange.Rows.Count + 2):
            val_kode = ws_dp.Cells(r, 4).Value
            if val_kode is not None and _clean(val_kode) == target_rup:
                found_row = r
                break

        if found_row is None:
            return False, f"Kode RUP '{kode_rup}' tidak ditemukan di sheet Daftar Paket."

        json_data = ws_dp.Cells(found_row, 11).Value
        if not json_data:
            return False, f"Snapshot untuk Kode RUP '{kode_rup}' kosong."

        # 2. Parse JSON
        fields = json.loads(str(json_data))

        # 3. Restore ke Master Data
        ws_master = wb.Sheets('Master Data')
        restored = 0
        for r in range(3, ws_master.UsedRange.Rows.Count + 3):
            lbl = ws_master.Cells(r, 1).Value
            if lbl is not None:
                lbl_str = str(lbl).strip()
                if lbl_str in fields:
                    ws_master.Cells(r, 2).Value = fields[lbl_str]
                    restored += 1
        scope = fields.get("Lingkup Pekerjaan", [])
        if isinstance(scope, list):
            for index, row_number in enumerate(range(14, 24)):
                ws_master.Cells(row_number, 5).Value = scope[index] if index < len(scope) else None

        return True, restored

    xl = None
    owned = False  # apakah kita yang spawn instance ini
    xl_running = None
    try:
        # Coba attach ke Excel yang sudah running dan punya file ini terbuka
        try:
            xl_running = win32com.client.GetObject(Class='Excel.Application')
            wb_found = None
            for i in range(1, xl_running.Workbooks.Count + 1):
                wb_i = xl_running.Workbooks(i)
                if _path_key(wb_i.FullName) == target_path_key:
                    wb_found = wb_i
                    break
            if wb_found is not None:
                # Attach mode: tulis langsung ke workbook yang terbuka
                ok, result = _do_restore(wb_found)
                if not ok:
                    return {'ok': False, 'error': result}
                wb_found.Save()
                # Refresh tampilan Excel
                xl_running.ScreenUpdating = True
                print(f"  Snapshot dimuat (attach). {result} field di-restore untuk RUP {kode_rup}")
                return {'ok': True}
        except Exception as attach_error:
            # Jangan fallback ke DispatchEx jika Excel sudah terdeteksi tetapi
            # COM/RPC sedang hang; itu akan membuka instance kedua dan ikut
            # mengunci workbook yang sedang dipakai user.
            if xl_running is not None:
                return {'ok': False, 'error': f'Excel COM/RPC tidak responsif: {attach_error}'}
            pass  # Excel memang tidak running — fallback ke DispatchEx

        # Fallback: buka instance baru
        owned = True
        xl = win32com.client.DispatchEx("Excel.Application")
        xl.Visible = False
        xl.DisplayAlerts = False
        wb = xl.Workbooks.Open(abs_path, ReadOnly=False, UpdateLinks=False)
        ok, result = _do_restore(wb)
        if not ok:
            wb.Close(SaveChanges=False)
            return {'ok': False, 'error': result}
        wb.Save()
        wb.Close(SaveChanges=False)
        print(f"  Snapshot dimuat. {result} field di-restore ke Master Data untuk RUP {kode_rup}")
        return {'ok': True}

    except Exception as e:
        return {'ok': False, 'error': str(e)}
    finally:
        if owned and xl:
            try: xl.Quit()
            except: pass
        pythoncom.CoUninitialize()


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════
def main():
    parser = argparse.ArgumentParser(description='Generate dokumen PPK PL')
    parser.add_argument('--mode', default='generate',
                        choices=['generate', 'pdf', 'pdf-only', 'selective', 'list-templates', 'paket_baru', 'multi',
                                 'commit-paket', 'nomor-baru', 'test-terbilang', 'muat-db', 'simpan-snapshot', 'muat-snapshot',
                                 'print-all', 'print-pilihan', 'self-check'])
    parser.add_argument('--kode', default='', help='Kode RUP untuk mode muat-db')
    parser.add_argument('--jenis', default='', help='Jenis surat (untuk nomor-baru)')
    parser.add_argument('--files', default='', help='Nomor urut template (1-based, pisah koma) untuk mode selective')
    parser.add_argument('--nomor', default='', help='Nomor urut paket (pisah koma) untuk mode print-pilihan. Contoh: 1,3,5')
    args = parser.parse_args()
    mode = args.mode

    if mode == 'self-check':
        print('=== Self-check PPK V2 ===')
        errors = []
        if not os.path.isfile(EXCEL_PATH):
            errors.append(f'Workbook tidak ditemukan: {EXCEL_PATH}')
        else:
            try:
                data = baca_excel(EXCEL_PATH)
                code, profile = _template_profile(data)
                templates = template_files_for_data(BASE_DIR, data)
                print(f'Workbook : {EXCEL_PATH}')
                print(f'Jenis    : {code} -> {profile}')
                print(f'Template : {len(templates)} file aktif')
                print(f'Tahap    : {normalize_document_stage(data.get("Tahap Dokumen"))}')
            except Exception as exc:
                errors.append(str(exc))
        if errors:
            for error in errors:
                print(f'ERROR: {error}')
            sys.exit(1)
        print('SELF-CHECK OK')
        sys.exit(0)

    # ── Mode list-templates ───────────────────────────────────────────────────
    if mode == 'list-templates':
        word_files = template_auto_detect(BASE_DIR)
        if not word_files:
            word_files = _word_files_fallback()
        for idx, (src_fname, out_prefix) in enumerate(word_files, 1):
            name = out_prefix if out_prefix else src_fname
            # bersihkan suffix .docx jika ada
            if name.lower().endswith('.docx'):
                name = name[:-5]
            print(f"{idx}. {name}")
        sys.exit(0)

    # ── Test terbilang ────────────────────────────────────────────────────────
    if mode == 'test-terbilang':
        print("=== Test Terbilang ===")
        tests = [
            (33300000,  'Tiga Puluh Tiga Juta Tiga Ratus Ribu Rupiah'),
            (1000000,   'Satu Juta Rupiah'),
            (1500,      'Seribu Lima Ratus Rupiah'),
            (100000000, 'Seratus Juta Rupiah'),
            (1100000,   'Satu Juta Seratus Ribu Rupiah'),
            (11000,     'Sebelas Ribu Rupiah'),
            (500000000, 'Lima Ratus Juta Rupiah'),
        ]
        all_ok = True
        for angka, expected in tests:
            hasil = terbilang_rupiah(angka)
            ok    = 'OK' if hasil == expected else 'FAIL'
            if hasil != expected:
                all_ok = False
            print(f"  [{ok}] {angka:>12,} -> {hasil}")
            if hasil != expected:
                print(f"        Expected: {expected}")
        print("\nTest Hari:")
        for n in [1, 10, 14, 30, 45, 90]:
            print(f"  {n:3d} hari -> {terbilang_hari(n)}")
        sys.exit(0 if all_ok else 1)

    # ── Print Pilihan PDF (generate paket tertentu + convert PDF) ────────────
    if mode == 'print-pilihan':
        if not args.nomor.strip():
            print("ERROR: --nomor diperlukan. Contoh: --nomor 1,3,5")
            sys.exit(1)
        try:
            nomor_filter = {int(x.strip()) for x in args.nomor.split(',') if x.strip()}
        except ValueError:
            print(f"ERROR: --nomor harus angka. Contoh: 1,3,5 (dapat: '{args.nomor}')")
            sys.exit(1)
        # Jalankan logika print-all dengan filter nomor
        args._print_pilihan_filter = nomor_filter
        mode = 'print-all'  # reuse blok print-all, filter diterapkan di sana

    # ── Print All PDF (generate semua paket + convert PDF) ───────────────────
    if mode == 'print-all':
        import openpyxl, win32com.client, pythoncom
        print("=== Mode: Print All PDF ===")

        # Baca Daftar Paket
        wb_ro = openpyxl.load_workbook(EXCEL_PATH, data_only=True, read_only=True)
        if 'Daftar Paket' not in wb_ro.sheetnames:
            wb_ro.close()
            print("ERROR: Sheet 'Daftar Paket' tidak ditemukan.")
            sys.exit(1)
        ws_dp = wb_ro['Daftar Paket']
        ws_md = wb_ro['Master Data']

        # Baca base_data dari Master Data
        base_data = {}
        for _row in ws_md.iter_rows(min_row=3, values_only=True):
            if _row[0] and not str(_row[0]).startswith('---') and not str(_row[0]).startswith('==='):
                base_data[str(_row[0])] = str(_row[1]) if _row[1] is not None else ''

        COL_FIELD = {
            3: 'Nama Paket (Lengkap)',
            4: 'Kode RUP',            5: 'Pagu Anggaran (Angka)',
            6: 'Nilai HPS (Angka)',   7: 'Lokasi Pekerjaan',
            8: 'Tanggal KAK & HPS',   9: 'Nama Penyedia',
        }
        paket_list = []
        rows_pa = list(ws_dp.iter_rows(min_row=2, values_only=True))
        wb_ro.close()

        for _i, _rv in enumerate(rows_pa):
            if not _rv[3]:  # Kode RUP kosong = baris kosong
                continue
            _ed = dict(base_data)
            for _ci, _fn in COL_FIELD.items():
                _v = _rv[_ci - 1]
                _ed[_fn] = str(_v) if _v is not None else ''
            # Nomor urut dari kolom 1 Daftar Paket (index 0) atau fallback index
            _nu = str(_rv[0]).strip() if _rv[0] else ''
            _ed['Nomor Urut Paket'] = _nu if _nu and _nu != 'None' else str(_i + 1)
            paket_list.append(_ed)

        if not paket_list:
            print("Tidak ada paket di sheet Daftar Paket.")
            sys.exit(0)

        # Filter nomor urut jika mode print-pilihan
        _pilihan_filter = getattr(args, '_print_pilihan_filter', None)
        if _pilihan_filter:
            paket_list = [_ed for _ed in paket_list
                          if _ed.get('Nomor Urut Paket', '').strip() in
                          {str(n) for n in _pilihan_filter}]
            if not paket_list:
                print(f"ERROR: Tidak ada paket dengan nomor {sorted(_pilihan_filter)}.")
                sys.exit(1)
            print(f"  Filter aktif: nomor {sorted(_pilihan_filter)}")

        print(f"  {len(paket_list)} paket ditemukan. Generate Word + PDF...")

        _base_pa   = '\\\\?\\' + BASE_DIR
        total_pdf  = 0

        pythoncom.CoInitialize()
        word = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0

            for _ed in paket_list:
                _ed = auto_terbilang(_ed)
                _kode_template, _folder_template = _template_profile(_ed)
                word_files = template_files_for_data(BASE_DIR, _ed)
                _nf  = _ed.get('Nama Paket (Lengkap)', '').strip()
                _nu  = _ed.get('Nomor Urut Paket', '').strip()
                # Strip prefix generik (sama seperti mode generate)
                _STRIP_PA = [
                    'Belanja Jasa Konsultansi Perencanaan Arsitektur-Jasa Arsitektur Lainnya ',
                    'Belanja Jasa Konsultansi Perencanaan Arsitektur-Jasa Arsitektur ',
                    'Belanja Jasa Konsultansi Perencanaan Rekayasa-Jasa Desain Rekayasa untuk Konstruksi ',
                    'Belanja Jasa Konsultansi Perencanaan ',
                    'Belanja Jasa Konsultansi ',
                    'Belanja Jasa ',
                ]
                _np = _nf
                for _pfx in _STRIP_PA:
                    if _np.startswith(_pfx):
                        _np = _np[len(_pfx):]
                        break
                _nf_folder = f"{_nu}. {_np}" if _nu and _nu != 'None' else _np
                safe_f = _re.sub(r'[\\/:*?"<>|]', '', _nf_folder).strip()
                out_dir_pa = os.path.join(_base_pa, safe_f or f'OUTPUT_{_i}')

                # Overwrite: hapus isi folder lama jika ada
                out_dir_real = os.path.join(BASE_DIR, safe_f or f'OUTPUT_{_i}')
                if os.path.isdir(out_dir_real):
                    for _f in os.listdir(out_dir_real):
                        try: os.remove(os.path.join(out_dir_real, _f))
                        except: pass
                os.makedirs(out_dir_pa, exist_ok=True)

                # Copy + proses Word
                for src_fname, out_prefix in word_files:
                    src = os.path.join(BASE_DIR, src_fname)
                    if not os.path.exists(src):
                        continue
                    dst_fname = f"{out_prefix}.docx" if out_prefix else src_fname
                    dst = os.path.join(out_dir_pa, dst_fname)
                    prepare_template_output(src, dst, EXCEL_PATH, _kode_template)
                    repl = build_replacements(_ed, src_fname, EXCEL_PATH)
                    process_docx(dst, repl)

                # Convert semua .docx → PDF, hapus .docx setelahnya
                docx_files = [f for f in os.listdir(out_dir_pa) if f.endswith('.docx')]
                n_ok = 0
                for _df in docx_files:
                    _dp = os.path.join(out_dir_pa, _df)
                    _pp = _dp[:-5] + '.pdf'
                    try:
                        _export_one(word, _dp, _pp)
                        os.remove(_dp)
                        n_ok += 1
                    except Exception as _ex:
                        print(f"    [WARN] PDF gagal ({_df}): {_ex}")
                total_pdf += n_ok
                print(f"  OK: {safe_f}/ — {n_ok} PDF")

        finally:
            if word:
                try: word.Quit(SaveChanges=False)
                except: pass
            pythoncom.CoUninitialize()

        _kill_lingering_word()
        print(f"\nSelesai. Total PDF: {total_pdf} dari {len(paket_list)} paket.")
        sys.exit(0)

    # ── Paket baru ────────────────────────────────────────────────────────────
    if mode == 'paket_baru':
        print("=== Mode: Paket Baru ===")
        paket_baru(EXCEL_PATH)
        print("Excel direset untuk paket baru.")
        sys.exit(0)

    # ── Nomor baru ────────────────────────────────────────────────────────────
    if mode == 'nomor-baru':
        if not args.jenis:
            print("ERROR: --jenis diperlukan. Pilihan:", list(_JENIS_TO_FIELD))
            sys.exit(1)
        print(f"=== Mode: Nomor Baru ({args.jenis}) ===")
        result = nomor_baru(EXCEL_PATH, args.jenis)
        if result['ok']:
            print(f"Nomor: {result['nomor']}")
            sys.exit(0)
        else:
            print(f"ERROR: {result['error']}")
            sys.exit(1)

    # ── Multi-generate ────────────────────────────────────────────────────────
    if mode == 'multi':
        print("=== Mode: Multi-Paket ===")
        generate_multi(EXCEL_PATH)
        sys.exit(0)

    # ── Commit paket ──────────────────────────────────────────────────────────
    if mode == 'commit-paket':
        print("=== Mode: Commit Paket ===")
        excel_data = baca_excel(EXCEL_PATH)
        excel_data = auto_terbilang(excel_data)
        commit_paket(EXCEL_PATH, excel_data)
        sys.exit(0)

    # ── Muat dari DB ──────────────────────────────────────────────────────────
    if mode == 'muat-db':
        print("=== Mode: Muat dari DB ===")
        kode_rup = args.kode.strip()
        if not kode_rup:
            print("ERROR: --kode diperlukan. Contoh: --kode 64312558")
            sys.exit(1)
        print(f"Mencari kode RUP: {kode_rup}...")
        data = tarik_dari_supabase(kode_rup)
        if not data:
            print(f"ERROR: Kode RUP '{kode_rup}' tidak ditemukan di database.")
            print("Pastikan sudah scrape via Streamlit or scraper_ppk.py --mode scrape")
            sys.exit(1)
        print(f"Ditemukan: {data.get('nama_paket', '-')}")
        isi_excel_dari_supabase(EXCEL_PATH, data)
        print(f"\nSelesai! Data paket dimuat ke Excel.")
        sys.exit(0)

    # ── Simpan Snapshot ───────────────────────────────────────────────────────
    if mode == 'simpan-snapshot':
        print("=== Mode: Simpan Snapshot ===")
        kode_rup = args.kode.strip()
        if not kode_rup:
            print("ERROR: --kode diperlukan. Contoh: --kode 64312558")
            sys.exit(1)
        result = simpan_snapshot(EXCEL_PATH, kode_rup)
        if result['ok']:
            print("Snapshot berhasil disimpan.")
            sys.exit(0)
        else:
            print(f"ERROR: {result['error']}")
            sys.exit(1)

    # ── Muat Snapshot ─────────────────────────────────────────────────────────
    if mode == 'muat-snapshot':
        print("=== Mode: Muat Snapshot ===")
        kode_rup = args.kode.strip()
        if not kode_rup:
            print("ERROR: --kode diperlukan. Contoh: --kode 64312558")
            sys.exit(1)
        result = muat_snapshot(EXCEL_PATH, kode_rup)
        if result['ok']:
            print("Snapshot berhasil dimuat.")
            sys.exit(0)
        else:
            print(f"ERROR: {result['error']}")
            sys.exit(1)

    # ── Generate & PDF ────────────────────────────────────────────────────────
    print(f"=== Mode: {'Generate Dokumen' if mode == 'generate' else 'Print PDF'} ===")

    # 1. Baca Excel
    excel_data = baca_excel(EXCEL_PATH)

    # 2. Auto-terbilang
    excel_data = auto_terbilang(excel_data)

    # 3. Tentukan folder output — strip prefix generik agar nama folder pendek + unik
    nama_lengkap = excel_data.get('Nama Paket (Lengkap)', '').strip()
    nomor_urut   = excel_data.get('Nomor Urut Paket', '').strip()
    # Hapus prefix panjang yang selalu sama agar nama folder tetap di bawah ~200 char
    _STRIP_PREFIX = [
        'Belanja Jasa Konsultansi Perencanaan Arsitektur-Jasa Arsitektur Lainnya ',
        'Belanja Jasa Konsultansi Perencanaan Arsitektur-Jasa Arsitektur ',
        'Belanja Jasa Konsultansi Perencanaan Rekayasa-Jasa Desain Rekayasa untuk Konstruksi ',
        'Belanja Jasa Konsultansi Perencanaan ',
        'Belanja Jasa Konsultansi ',
        'Belanja Jasa ',
    ]
    nama_pendek = nama_lengkap
    for _pfx in _STRIP_PREFIX:
        if nama_pendek.startswith(_pfx):
            nama_pendek = nama_pendek[len(_pfx):]
            break
    nama_folder  = f"{nomor_urut}. {nama_pendek}" if nomor_urut and nomor_urut != 'None' else nama_pendek
    safe_nama    = _re.sub(r'[\\/:*?"<>|]', '', nama_pendek).strip() if nama_pendek else ''
    safe_folder  = _re.sub(r'[\\/:*?"<>|]', '', nama_folder).strip() if nama_folder else safe_nama
    _base        = '\\\\?\\' + BASE_DIR
    out_dir      = os.path.join(_base, safe_folder) if safe_folder \
                   else os.path.join(_base, f'OUTPUT_{datetime.now().strftime("%Y%m%d_%H%M")}')

    # ── Mode PDF (konversi docx existing) ────────────────────────────────────
    if mode == 'pdf':
        if not os.path.isdir(out_dir):
            print(f"ERROR: Folder output tidak ditemukan: {out_dir}")
            print("Jalankan generate dulu sebelum print PDF.")
            sys.exit(1)
        errors = print_semua_pdf(out_dir)
        n_pdf  = len([f for f in os.listdir(out_dir) if f.endswith('.pdf')])
        log    = {
            'waktu':         datetime.now().strftime('%Y-%m-%d %H:%M'),
            'folder':        os.path.basename(out_dir),
            'total_replace': '-',
            'field_kosong':  [],
            'status':        f"PDF: {n_pdf} file" + (f" ({len(errors)} error)" if errors else ''),
        }
        tulis_log_excel(EXCEL_PATH, log)
        sys.exit(1 if errors else 0)

    # ── Mode Generate ─────────────────────────────────────────────────────────
    # 4. Validasi wajib
    kosong = validasi_wajib(excel_data)
    if kosong:
        print(f"VALIDASI GAGAL — {len(kosong)} field wajib belum diisi:")
        for k in kosong:
            print(f"  - {k}")
        tulis_log_excel(EXCEL_PATH, {
            'waktu':         datetime.now().strftime('%Y-%m-%d %H:%M'),
            'folder':        '-',
            'total_replace': 0,
            'field_kosong':  kosong,
            'status':        f'GAGAL — {len(kosong)} field wajib kosong',
        })
        sys.exit(1)

    stage = normalize_document_stage(excel_data.get("Tahap Dokumen"))
    stage_missing = validate_document_stage(excel_data)
    if stage_missing:
        print(f"VALIDASI TAHAP {stage} GAGAL — field final belum lengkap:")
        for field in stage_missing:
            print(f"  - {field}")
        sys.exit(1)
    print(f"  Tahap dokumen: {stage}")

    # 5. Pilih profil template dari satu kode di Master Data
    try:
        _kode_template, _folder_template = _template_profile(excel_data)
        word_files = template_files_for_data(BASE_DIR, excel_data)
        print(f"  Profil template: {_kode_template} ({_folder_template}) — {len(word_files)} template")
    except (ValueError, FileNotFoundError) as exc:
        print(f"ERROR TEMPLATE: {exc}")
        sys.exit(1)

    # Filter files jika mode selective
    if mode == 'selective':
        if not args.files:
            print("ERROR: --files diperlukan untuk mode selective. Contoh: --files 1,3")
            sys.exit(1)
        try:
            indices = [int(i.strip()) - 1 for i in args.files.split(',') if i.strip()]
            word_files = [word_files[idx] for idx in indices if 0 <= idx < len(word_files)]
        except Exception as e:
            print(f"ERROR parsing --files: {e}")
            sys.exit(1)
        if not word_files:
            print("ERROR: Tidak ada template yang cocok dengan filter --files.")
            sys.exit(1)

    # 7. Buat folder output
    os.makedirs(out_dir, exist_ok=True)
    print(f"  Folder: {os.path.basename(out_dir)}/")

    # 8. Copy + process .docx
    filled    = 0
    all_empty = [k for k in FIELD_MAP if not excel_data.get(k) or excel_data[k] in ('', 'None')]
    pdf_only  = (mode == 'pdf-only' or mode == 'selective')
    # dst_list: list of (dst_path, dst_fname, n_replace) untuk batch PDF
    dst_list  = []
    for src_fname, out_prefix in word_files:
        src = os.path.join(BASE_DIR, src_fname)
        if not os.path.exists(src):
            print(f"  SKIP (tidak ada): {src_fname}")
            continue
        repl = build_replacements(excel_data, src_fname)
        dst_fname = f"{out_prefix}.docx" if out_prefix else os.path.basename(src_fname)
        dst       = os.path.join(out_dir, dst_fname)
        _rows, _resource_tables = prepare_template_output(
            src, dst, EXCEL_PATH, _kode_template
        )
        if _rows:
            print(f"  Personil Excel -> Word: {_rows} baris")
        if _resource_tables:
            print(f"  Sumber daya PK Excel -> KAK: {_resource_tables} tabel")
        n = process_docx(dst, repl)
        filled += n
        if pdf_only:
            dst_list.append((dst, dst_fname, n))
        else:
            print(f"  OK ({n} replace): {dst_fname}")

    # 8b. Batch PDF — satu instance Word untuk semua file
    if pdf_only and dst_list:
        import win32com.client, pythoncom
        pythoncom.CoInitialize()
        word = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            for dst, dst_fname, n in dst_list:
                pdf_path = dst[:-5] + '.pdf'
                try:
                    _export_one(word, dst, pdf_path)
                    os.remove(dst)
                    print(f"  OK ({n} replace): {dst_fname[:-5]}.pdf")
                except Exception as e:
                    print(f"  PDF GAGAL ({dst_fname}): {e}")
        finally:
            if word:
                try: word.Quit(SaveChanges=False)
                except: pass
            pythoncom.CoUninitialize()

    print(f"\nOutput : {out_dir}")
    print(f"Replace: {filled} total")
    if all_empty:
        print(f"Field belum diisi ({len(all_empty)}):")
        for f in all_empty:
            print(f"  - {f}")

    if pdf_only:
        n_pdf = len([f for f in os.listdir(out_dir) if f.endswith('.pdf')])
        status_msg = f'SUKSES {mode.upper()} ({filled} replace, {n_pdf} PDF)'
    else:
        # 8b. Cetak PDF otomatis setelah generate Word
        print("\nMenjalankan cetak PDF otomatis...")
        errors_pdf = print_semua_pdf(out_dir)
        n_pdf = len([f for f in os.listdir(out_dir) if f.endswith('.pdf')])
        status_msg = f'SUKSES ({filled} replace, {n_pdf} PDF)'
        if errors_pdf:
            status_msg += f' - {len(errors_pdf)} PDF gagal'

    # 9. Tulis log ke Excel
    tulis_log_excel(EXCEL_PATH, {
        'waktu':         datetime.now().strftime('%Y-%m-%d %H:%M'),
        'folder':        os.path.basename(out_dir),
        'total_replace': filled,
        'field_kosong':  all_empty,
        'status':        status_msg,
    })

    # 10. Catat riwayat
    catat_riwayat(EXCEL_PATH, excel_data, out_dir, status_msg)

    # 11. Sync Supabase (non-fatal)
    sb_result = sync_supabase(excel_data)
    if sb_result['ok']:
        print("  Supabase: tersimpan di draft_paket_pl")
    else:
        print(f"  Supabase: {sb_result['error']} (non-fatal)")

    _kill_lingering_word()
    print("\nSelesai!")


if __name__ == '__main__':
    main()
