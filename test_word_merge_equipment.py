import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import Workbook

import word_merge


class DokpilEquipmentTests(unittest.TestCase):
    def test_reader_accepts_tender_sheet_and_skips_header_empty_rows(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "paket.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.title = "Alat & Personil"
            ws.append([])
            ws.append([])
            ws.append([])
            ws.append([None, "No.", "Jenis Alat", "Kapasitas", "Jumlah"])
            ws.append([None, 1, "Concrete Mixer", "0.3-0.6 M3", "1 Unit"])
            ws.append([None, 2, "Concrete Vibrator", "5 HP", "1 Unit"])
            ws.append([None, 3, 0, 0, 0])
            wb.save(path)

            self.assertEqual(
                word_merge._read_dokpil_equipment(str(path)),
                [
                    {"no": "1", "jenis": "Concrete Mixer", "kapasitas": "0.3-0.6 M3", "jumlah": "1 Unit"},
                    {"no": "2", "jenis": "Concrete Vibrator", "kapasitas": "5 HP", "jumlah": "1 Unit"},
                ],
            )

    def test_tender_six_and_eight_column_tables_are_filled(self):
        equipment = [
            {"no": "1", "jenis": "Concrete Mixer", "kapasitas": "0.3-0.6 M3", "jumlah": "1 Unit"},
            {"no": "2", "jenis": "Concrete Vibrator", "kapasitas": "5 HP", "jumlah": "1 Unit"},
        ]
        data = {"_dokpil_equipment": equipment}
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "dokpil.docx"
            doc = Document()
            for headers in (
                ["No", "Jenis", "Merek dan Tipe*)", "Kapasitas", "Jumlah", "Kepemilikan/status"],
                ["No", "Nama Peralatan Utama*)", "Merk dan Tipe**)", "Kapasitas**)", "Jumlah**)", "Kondisi**)", "Status Kepemilikan**)", "Keterangan"],
            ):
                table = doc.add_table(rows=2, cols=len(headers))
                for cell, value in zip(table.rows[0].cells, headers):
                    cell.text = value
                for cell in table.rows[1].cells:
                    cell.text = "..."
            doc.save(path)

            word_merge._prepare_dokpil_equipment_docx(str(path), data)
            merged = Document(path)
            self.assertEqual(len(merged.tables), 2)
            for table in merged.tables:
                self.assertEqual(len(table.rows), 3)
                self.assertEqual([table.rows[i].cells[0].text.strip() for i in (1, 2)], ["1", "2"])
                self.assertEqual(
                    [table.rows[i].cells[1].text.strip() for i in (1, 2)],
                    ["Concrete Mixer", "Concrete Vibrator"],
                )


if __name__ == "__main__":
    unittest.main()
