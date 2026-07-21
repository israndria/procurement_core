"""
parse_reviu.py — Parse Dokumen Tender (Single PDF or Directory) → isi database_reviu + database_dokpil
====================================================================================================
Optimasi untuk Pokja 030:
1. Mendukung scan folder untuk file spesifik (Daftar Peralatan, Personil, Kuantitas, dll).
2. Perbaikan regex dan table extractor untuk struktur tabel yang terpecah (multi-line cells).
3. Penanganan "Resiko Tertinggi" pada RK3K yang lebih presisi.
"""

import sys
import os
import json
import re
import concurrent.futures as _cf
from bs4 import BeautifulSoup

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ─── Load env Supabase ───────────────────────────────────────────────────────
def _load_env():
    secret_root = os.path.normpath(os.environ.get(
        "POKJA_SECRET_ROOT", os.path.join(os.path.dirname(BASE_DIR), "Secrets")
    ))
    env_file = os.path.join(secret_root, "secret_supabase.env")
    if os.path.exists(env_file):
        with open(env_file, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if "=" in line and not line.startswith("#"):
                    k, v = line.split("=", 1)
                    os.environ.setdefault(k.strip(), v.strip().strip('"').strip("'"))

_load_env()
SUPABASE_URL = os.environ.get("SUPABASE_URL", "")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "")

# ─── Fetch cara_pembayaran dari Supabase (cache lokal TTL 1 jam) ─────────────
_CP_CACHE_FILE = os.path.join(BASE_DIR, "_cara_pembayaran_cache.json")

def fetch_cara_pembayaran():
    import time as _time
    # Pakai cache lokal jika ada dan belum expire (TTL 3600s)
    try:
        if os.path.exists(_CP_CACHE_FILE):
            _age = _time.time() - os.path.getmtime(_CP_CACHE_FILE)
            if _age < 3600:
                with open(_CP_CACHE_FILE, encoding="utf-8") as _fc:
                    return json.load(_fc)
    except Exception:
        pass
    try:
        import httpx
        r = httpx.get(
            f"{SUPABASE_URL}/rest/v1/cara_pembayaran?select=id,keyword,teks",
            headers={"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}"},
            timeout=10
        )
        if r.status_code == 200:
            data = r.json()
            try:
                with open(_CP_CACHE_FILE, "w", encoding="utf-8") as _fw:
                    json.dump(data, _fw, ensure_ascii=False)
            except Exception:
                pass
            return data
    except Exception:
        pass
    return [
        {
            "id": "monthly_certificate",
            "keyword": ["monthly", "monthly certificate", "bulanan"],
            "teks": "(Monthly Certificate) Pembayaran dilakukan dengan cara didasarkan pada hasil pengukuran bersama atas pekerjaan yang benar-benar telah dilaksanakan secara bulanan"
        },
        {
            "id": "termin",
            "keyword": ["termin", "pembayaran termin", "angsuran"],
            "teks": "(Termin) Pembayaran dilakukan dengan cara pembayarannya didasarkan pada hasil pengukuran bersama atas pekerjaan yang benar-benar telah dilaksanakan secara cara angsuran"
        }
    ]

def deteksi_cara_pembayaran(daftar_cp, bidang=""):
    bidang_lower = (bidang or "").lower().strip()
    KEYWORD_BINAMARGA = ["bina marga", "jalan", "jembatan", "perkerasan", "rigid", "flexible", "hotmix", "aspal", "overlay"]
    is_binamarga = any(kw in bidang_lower for kw in KEYWORD_BINAMARGA)
    id_pilih = "monthly_certificate" if is_binamarga else "termin"
    for cp in daftar_cp:
        if cp["id"] == id_pilih:
            # Bersihkan teks cara pembayaran dari redundansi
            teks = cp["teks"]
            teks = re.sub(r"\bsecara cara\b", "secara", teks)
            teks = re.sub(r"\s{2,}", " ", teks).strip()
            return cp["id"], teks
    return daftar_cp[0]["id"], daftar_cp[0]["teks"]

# ─── Helpers ────────────────────────────────────────────────────────────
def bersihkan(s):
    if not s: return ""
    s = str(s)
    # Fix encoding rusak: bytes UTF-8 dibaca sebagai latin-1/cp1252 (â€" → –)
    try:
        s = s.encode("latin-1").decode("utf-8")
    except (UnicodeDecodeError, UnicodeEncodeError):
        pass
    # Fix � (replacement char) yang muncul di antara angka → en dash
    s = re.sub(r"(?<=[\d,])\s*�+\s*(?=[\d,])", " – ", s)
    # Hapus replacement char yang tersisa
    s = s.replace("�", "")
    # Fix PDF hyphenation: "kata- kata" -> "katakara"
    import re as _re
    s = _re.sub(r"(\w+)-\s+(\w)", lambda m: m.group(1) + m.group(2), s)
    return " ".join(s.split()).strip()

def normalisasi_pengalaman(s):
    if not s: return 0
    s = str(s).lower()
    if "-" in s or "0" in s or not any(c.isdigit() for c in s): return 0
    m = re.search(r"(\d+)", s)
    return int(m.group(1)) if m else 0

def extract_rekap_from_xlsx(path):
    try:
        import pandas as pd
        xl = pd.ExcelFile(path)
        sheet_name = ""
        for s in xl.sheet_names:
            if any(kw in s.upper() for kw in ["REKAP", "BOQ", "DAFTAR", "KUANTITAS"]):
                sheet_name = s
                break
        if not sheet_name: sheet_name = xl.sheet_names[0]
        
        df = pd.read_excel(path, sheet_name=sheet_name, header=None)
        items = []
        for index, row in df.iterrows():
            # Gabungkan semua kolom dalam baris ini untuk mencari kata "DIVISI"
            full_row_text = " ".join([str(v) for v in row if v and not pd.isna(v)])
            
            # Cari pola "DIVISI [X] [DESKRIPSI]"
            m_div = re.search(r"(DIVISI\s+[\d\w]+)\s*(.*)", full_row_text, re.IGNORECASE)
            if m_div:
                div_num = m_div.group(1).upper()
                div_desc = bersihkan(m_div.group(2))
                
                # Filter jika ternyata ini header (misal: DIVISI URAIAN)
                if any(kw in div_desc.upper() for kw in ["URAIAN", "HARGA", "BOBOT", "KETERANGAN"]):
                    continue
                
                # Bersihkan sisa-sisa karakter sampah di akhir
                div_desc = re.sub(r"[\d.,\s-]+$", "", div_desc).strip()
                
                full_val = f"{div_num} {div_desc}".strip()
                if len(full_val) > 8 and full_val not in items:
                    items.append(full_val)
            
            # Fallback untuk pola Romawi
            elif re.match(r"^[ \t]*([IVXLC]+\.?)[ \t]+([A-Z][A-Z \/,()&]{5,100})", full_row_text):
                items.append(bersihkan(full_row_text))
                
        return items
    except Exception: return []

def extract_text_from_docx(path):
    try:
        from docx import Document
        doc = Document(path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return "\n".join(full_text)
    except Exception: return ""

def extract_tables_from_docx(path):
    try:
        from docx import Document
        doc = Document(path)
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append(rows)
        return tables
    except Exception: return []

def find_files_by_keywords(directory, keywords):
    """Mencari file dengan prioritas format: xlsx > docx > pdf."""
    found = []
    if not os.path.isdir(directory): return found
    
    all_matching = []
    for root, _dirs, files in os.walk(directory):
      for name in files:
        ext = name.lower().split(".")[-1]
        if ext not in ["pdf", "docx", "xlsx", "xlsm"]: continue
        if any(kw.lower() in name.lower() for kw in keywords):
            priority = 0
            if ext in ["xlsx", "xlsm"]: priority = 3
            elif ext == "docx": priority = 2
            elif ext == "pdf": priority = 1
            all_matching.append((priority, os.path.join(root, name)))
            
    # Sort berdasarkan priority DESC
    all_matching.sort(key=lambda x: x[0], reverse=True)
    return [x[1] for x in all_matching]

# ─── Logic Parsers Per Bagian ───────────────────────────────────────────────

def _parse_fung_bangunan(pdf_or_path, hasil):
    """E2: Fungsi Bangunan — cari di KAK/DOCX, bukan Draft PDF"""
    fname = os.path.basename(pdf_or_path if isinstance(pdf_or_path, str) else "").lower()
    if "gambar" in fname or "draft_pokja" in fname: return False

    if isinstance(pdf_or_path, str):
        if pdf_or_path.lower().endswith(".docx"):
            teks_list = [extract_text_from_docx(pdf_or_path)]
        else: return False
    else:
        teks_list = [p.extract_text() or "" for p in pdf_or_path.pages]

    # Kata yang menunjukkan ini BUKAN nama fungsi bangunan (pihak, koordinator, dll)
    FORBIDDEN_KEYWORDS = [
        "digambar oleh", "diperiksa oleh", "no. gambar", "judul gambar", "skala :", "proyek :",
        "koordinator", "rencana kerja", "konsultan", "pemberi tugas", "pengawas", "cv ", "pt ",
        "dinas ", "unit kerja", "satuan kerja", "upt ", "badan ", "kepala "
    ]
    # Kata yang HARUS ada agar match dianggap fungsi bangunan
    VALID_KEYWORDS = [
        "gedung", "bangunan", "jalan", "jembatan", "irigasi", "saluran", "drainase",
        "instalasi", "jaringan", "konstruksi", "infrastruktur", "fasilitas", "rehabilitasi",
        "renovasi", "pemeliharaan", "peningkatan", "pembangunan"
    ]

    for teks in teks_list:
        # Pola 1: "Nama Pekerjaan :" atau "Pekerjaan :"
        m_pek = re.search(r"(?:Nama\s+)?Pekerjaan\s*[:/]\s*([^\n]{10,200})", teks, re.IGNORECASE)
        if m_pek:
            val = bersihkan(m_pek.group(1))
            val_lower = val.lower()
            if (10 < len(val) < 180
                    and not any(kw in val_lower for kw in FORBIDDEN_KEYWORDS)
                    and any(kw in val_lower for kw in VALID_KEYWORDS)):
                hasil["reviu"]["E2"]["nilai"] = val
                hasil["reviu"]["E2"]["status"] = "terisi"
                return True

        # Pola 2: "Fungsi Bangunan :" atau "Jenis Bangunan :"
        m_fung = re.search(r"(?:Fungsi|Jenis)\s+Bangunan\s*[:/]\s*([^\n]{5,100})", teks, re.IGNORECASE)
        if m_fung:
            val = bersihkan(m_fung.group(1))
            if 5 < len(val) < 100:
                hasil["reviu"]["E2"]["nilai"] = val
                hasil["reviu"]["E2"]["status"] = "terisi"
                return True

        # Pola 3: Judul dokumen / nama pekerjaan dari header KAK/RKS
        # Contoh: "PEMBANGUAN BANGUNAN CYTOTOXIC - RSUD DATU SANGGUL KABUPATEN TAPIN"
        # Untuk E2, ambil bagian nama pekerjaan saja (sebelum " - " jika ada)
        m_judul = re.search(
            r"(?:KERANGKA ACUAN KERJA|SPESIFIKASI TEKNIS|RKS)\s*\n+([^\n]{10,150})",
            teks, re.IGNORECASE
        )
        if m_judul:
            val_full = bersihkan(m_judul.group(1))
            # Potong bagian nama instansi setelah " - " atau " – "
            val = re.split(r"\s+[-–]\s+", val_full)[0].strip()
            if len(val) < 10:
                val = val_full  # fallback jika terlalu pendek setelah split
            val_lower = val.lower()
            if (10 < len(val) < 150
                    and not any(kw in val_lower for kw in FORBIDDEN_KEYWORDS)
                    and any(kw in val_lower for kw in VALID_KEYWORDS)):
                hasil["reviu"]["E2"]["nilai"] = val
                hasil["reviu"]["E2"]["status"] = "terisi"
                return True

        # Pola 4: Tujuan/Maksud — ambil baris pertama saja, maksimum 80 karakter
        m = re.search(r"(?:Tujuan|Maksud)\s+dari\s+(?:pelaksanaan\s+)?(?:pekerjaan\s+)?([^\n]{10,80})", teks, re.IGNORECASE)
        if not m:
            m = re.search(r"(?:Tujuan|Maksud)\s*[:/]?\s*\n*([^\n]{20,80})", teks, re.IGNORECASE)
        if m:
            val = bersihkan(m.group(1))
            # Buang awalan yang tidak informatif
            val = re.sub(r"^(?:adalah|ini|untuk|agar|supaya)\s+", "", val, flags=re.IGNORECASE).strip()
            val_lower = val.lower()
            if (10 < len(val) < 80
                    and not any(kw in val_lower for kw in FORBIDDEN_KEYWORDS)
                    and any(kw in val_lower for kw in VALID_KEYWORDS)):
                hasil["reviu"]["E2"]["nilai"] = val
                hasil["reviu"]["E2"]["status"] = "terisi"
                return True
    return False

def _parse_peralatan(pdf_or_path, hasil):
    """E9-E26: Peralatan"""
    alat_list = []
    if isinstance(pdf_or_path, str):
        if pdf_or_path.lower().endswith(".docx"):
            tables = extract_tables_from_docx(pdf_or_path)
        else: return False
    else:
        tables = []
        for page in pdf_or_path.pages:
            t = page.extract_tables()
            if t: tables.extend(t)

    for table in tables:
        if not table or len(table) < 2: continue
        
        # Deteksi index kolom
        idx_nama = -1; idx_kap = -1; idx_jml = -1
        header = [str(c or "").lower() for c in table[0]]
        for i, c in enumerate(header):
            if any(kw in c for kw in ["jenis", "nama", "peralatan"]): idx_nama = i
            if any(kw in c for kw in ["kapasitas", "ukuran"]): idx_kap = i
            if any(kw in c for kw in ["jumlah", "qty"]): idx_jml = i
        
        if idx_nama == -1: continue # Bukan tabel peralatan
        
        for row in table[1:]:
            if len(row) <= idx_nama: continue
            nama_raw = str(row[idx_nama] or "").strip()
            if not nama_raw or nama_raw.isdigit() or len(nama_raw) < 3: continue
            if any(kw in nama_raw.lower() for kw in ["no", "jenis", "peralatan", "nama"]): continue

            # Gabungkan multiline dalam satu cell menjadi satu nama alat (pisahkan dengan spasi)
            # Contoh: "Alat pancang mini\npile." → "Alat pancang mini pile."
            nama_joined = " ".join(v.strip() for v in nama_raw.split("\n") if v.strip())
            nama_joined = re.sub(r"\s+", " ", nama_joined).strip()

            kap_raw = str(row[idx_kap] or "").strip() if idx_kap != -1 else ""
            kap_joined = " ".join(v.strip() for v in kap_raw.split("\n") if v.strip())
            kap_joined = re.sub(r"\s+", " ", kap_joined).strip()

            jml_raw = str(row[idx_jml] or "").strip() if idx_jml != -1 else "1"
            jml_joined = " ".join(v.strip() for v in jml_raw.split("\n") if v.strip())
            jml_joined = re.sub(r"\s+", " ", jml_joined).strip() or "1"

            n = bersihkan(nama_joined)
            if not n or n.isdigit() or any(kw in n.lower() for kw in ["no", "jenis", "peralatan", "nama"]):
                continue
            k = bersihkan(kap_joined)
            j = bersihkan(jml_joined)
            if not any(c.isdigit() for c in j): j = "1 Unit"
            elif "unit" not in j.lower() and "set" not in j.lower(): j += " Unit"
            alat_list.append({"nama": n, "kapasitas": k, "jumlah": j})
    
    if not alat_list: return False

    # Isi ke hasil
    ALAT_CELLS  = ["E9","E10","E11","E12","E13","E14"]
    KAP_CELLS   = ["E15","E16","E17","E18","E19","E20"]
    JML_CELLS   = ["E21","E22","E23","E24","E25","E26"]
    for i, alat in enumerate(alat_list[:6]):
        hasil["reviu"][ALAT_CELLS[i]]["nilai"] = alat["nama"]
        hasil["reviu"][ALAT_CELLS[i]]["status"] = "terisi"
        hasil["reviu"][KAP_CELLS[i]]["nilai"] = alat["kapasitas"]
        hasil["reviu"][KAP_CELLS[i]]["status"] = "terisi"
        hasil["reviu"][JML_CELLS[i]]["nilai"] = alat["jumlah"]
        hasil["reviu"][JML_CELLS[i]]["status"] = "terisi"
    for i in range(len(alat_list), 6):
        hasil["reviu"][ALAT_CELLS[i]]["status"] = "tidak_ada"
        hasil["reviu"][KAP_CELLS[i]]["status"] = "tidak_ada"
        hasil["reviu"][JML_CELLS[i]]["status"] = "tidak_ada"
    return True

def _parse_peralatan(pdf_or_path, hasil):
    """E9-E26: Peralatan. Handles SPSE rows with multiple newline-separated tools."""
    alat_list = []
    if isinstance(pdf_or_path, str):
        if pdf_or_path.lower().endswith(".docx"):
            tables = extract_tables_from_docx(pdf_or_path)
        else:
            return False
    else:
        tables = []
        for page in pdf_or_path.pages:
            t = page.extract_tables()
            if t:
                tables.extend(t)

    for table in tables:
        if not table or len(table) < 2:
            continue
        idx_nama = -1; idx_kap = -1; idx_jml = -1
        header = [str(c or "").lower() for c in table[0]]
        for i, c in enumerate(header):
            if any(kw in c for kw in ["jenis", "nama", "peralatan"]): idx_nama = i
            if any(kw in c for kw in ["kapasitas", "ukuran"]): idx_kap = i
            if any(kw in c for kw in ["jumlah", "qty"]): idx_jml = i
        if idx_nama == -1:
            continue

        for row in table[1:]:
            if len(row) <= idx_nama:
                continue
            nama_raw = str(row[idx_nama] or "").strip()
            if not nama_raw or nama_raw.isdigit() or len(nama_raw) < 3:
                continue
            if any(kw in nama_raw.lower() for kw in ["no", "jenis", "peralatan", "nama"]):
                continue
            kap_raw = str(row[idx_kap] or "").strip() if idx_kap != -1 else ""
            jml_raw = str(row[idx_jml] or "").strip() if idx_jml != -1 else "1"
            nama_parts = [v.strip() for v in nama_raw.split("\n") if v.strip()]
            kap_parts = [v.strip() for v in kap_raw.split("\n") if v.strip()]
            jml_parts = [v.strip() for v in jml_raw.split("\n") if v.strip()]
            n_items = len(nama_parts) if len(nama_parts) > 1 and len(jml_parts) in (0, len(nama_parts)) and len(kap_parts) in (0, len(nama_parts)) else 1
            for i_part in range(n_items):
                nama_joined = nama_parts[i_part] if n_items > 1 else " ".join(nama_parts)
                kap_joined = kap_parts[i_part] if n_items > 1 and i_part < len(kap_parts) else " ".join(kap_parts)
                jml_joined = jml_parts[i_part] if n_items > 1 and i_part < len(jml_parts) else (" ".join(jml_parts) or "1")
                n = bersihkan(re.sub(r"\s+", " ", nama_joined).strip())
                if not n or n.isdigit() or any(kw in n.lower() for kw in ["no", "jenis", "peralatan", "nama"]):
                    continue
                k = bersihkan(re.sub(r"\s+", " ", kap_joined).strip())
                j = bersihkan(re.sub(r"\s+", " ", jml_joined).strip())
                if not any(c.isdigit() for c in j): j = "1 Unit"
                elif "unit" not in j.lower() and "set" not in j.lower(): j += " Unit"
                alat_list.append({"nama": n, "kapasitas": k, "jumlah": j})

    if not alat_list:
        return False
    ALAT_CELLS  = ["E9","E10","E11","E12","E13","E14"]
    KAP_CELLS   = ["E15","E16","E17","E18","E19","E20"]
    JML_CELLS   = ["E21","E22","E23","E24","E25","E26"]
    for i, alat in enumerate(alat_list[:6]):
        hasil["reviu"][ALAT_CELLS[i]]["nilai"] = alat["nama"]
        hasil["reviu"][ALAT_CELLS[i]]["status"] = "terisi"
        hasil["reviu"][KAP_CELLS[i]]["nilai"] = alat["kapasitas"]
        hasil["reviu"][KAP_CELLS[i]]["status"] = "terisi"
        hasil["reviu"][JML_CELLS[i]]["nilai"] = alat["jumlah"]
        hasil["reviu"][JML_CELLS[i]]["status"] = "terisi"
    for i in range(len(alat_list), 6):
        hasil["reviu"][ALAT_CELLS[i]]["status"] = "tidak_ada"
        hasil["reviu"][KAP_CELLS[i]]["status"] = "tidak_ada"
        hasil["reviu"][JML_CELLS[i]]["status"] = "tidak_ada"
    return True

def _fallback_sbu_dari_nama(nama_tender: str):
    n = (nama_tender or "").lower()
    rules = [
        (["normalisasi", "sungai", "sei.", " sei ", "irigasi", "drainase", "saluran"], "Subklasifikasi BS010 (KBLI 2020) Konstruksi Bangunan Prasarana Sumber Daya Air", "Subklasifikasi SI001 (KBLI 2015) Jasa Pelaksana Konstruksi Saluran Air, Pelabuhan, Dam, dan Prasarana Sumber Daya Air Lainnya"),
        (["jalan", "pengaspalan", "aspal", "hotmix", "rabat beton"], "Subklasifikasi BS001 (KBLI 2020) Konstruksi Bangunan Sipil Jalan", "Subklasifikasi SI003 (KBLI 2015) Jasa Pelaksana Konstruksi Jalan Raya (kecuali Jalan Layang), Jalan, Rel Kereta Api, dan Landas Pacu Bandara"),
        (["jembatan", "plat duiker", "gorong"], "Subklasifikasi BS002 (KBLI 2020) Bangunan Sipil Jembatan, Jalan Layang, Fly Over, dan Underpass", "Subklasifikasi SI004 (KBLI 2015) Jasa Pelaksana Konstruksi Jembatan, Jalan Layang, Terowongan, dan Subways"),
        (["gedung", "bangunan"], "Subklasifikasi BG009 (KBLI 2020) Konstruksi Gedung Lainnya", "Subklasifikasi BG009 (KBLI 2015) Jasa Pelaksana Konstruksi Bangunan Gedung Lainnya"),
    ]
    for kws, baru, lama in rules:
        if any(k in n for k in kws):
            return baru, lama
    return "", ""

def _parse_personil(pdf_or_path, hasil):
    """E27-E32: Personil"""
    personil_list = []
    if isinstance(pdf_or_path, str):
        if pdf_or_path.lower().endswith(".docx"):
            tables = extract_tables_from_docx(pdf_or_path)
        else: return False
    else:
        tables = []
        for page in pdf_or_path.pages:
            t = page.extract_tables()
            if t: tables.extend(t)

    for table in tables:
        if not table or len(table) < 2: continue
        idx_jab = -1; idx_exp = -1; idx_sert = -1
        header = [str(c or "").lower() for c in table[0]]
        for i, c in enumerate(header):
            if "jabatan" in c: idx_jab = i
            if "pengalaman" in c: idx_exp = i
            if any(kw in c for kw in ["keterangan", "sertifikat", "skk", "skt", "bukti", "syarat"]): idx_sert = i
        
        if idx_jab == -1: continue
        
        for row in table[1:]:
            if len(row) <= idx_jab: continue
            jab_raw = str(row[idx_jab] or "").strip()
            if not jab_raw or len(jab_raw) < 3 or any(kw in jab_raw.lower() for kw in ["no", "jabatan"]): continue
            
            jabs = [j.strip() for j in jab_raw.split("\n") if len(j.strip()) > 3]
            exps_raw = str(row[idx_exp] or "").split("\n") if idx_exp != -1 else ["0"]
            srts_raw = str(row[idx_sert] or "").split("\n") if idx_sert != -1 else [""]
            
            for i, jb in enumerate(jabs):
                ex_val = exps_raw[i] if i < len(exps_raw) else exps_raw[-1]
                ex = normalisasi_pengalaman(ex_val)
                step = len(srts_raw) // len(jabs) if len(jabs) > 0 else 1
                st = " ".join(srts_raw[i*step : (i+1)*step]) if step > 0 else " ".join(srts_raw)
                personil_list.append({"jabatan": bersihkan(jb), "pengalaman": ex, "sertifikat": bersihkan(st)})
    
    if not personil_list: return False

    p_teknis = None; p_k3 = None
    for p in personil_list:
        if any(kw in p["jabatan"].lower() for kw in ["k3", "keselamatan", "hse"]):
            if not p_k3: p_k3 = p
        else:
            if not p_teknis: p_teknis = p
            
    if p_teknis:
        hasil["reviu"]["E27"]["nilai"] = "Pelaksana Lapangan"
        hasil["reviu"]["E27"]["status"] = "terisi"
        hasil["reviu"]["E28"]["nilai"] = p_teknis["pengalaman"]
        hasil["reviu"]["E28"]["status"] = "terisi"
        skt = p_teknis["sertifikat"]
        hasil["reviu"]["E29"]["nilai"] = skt
        hasil["reviu"]["E29"]["status"] = "terisi" if skt else "tidak_ada"
    if p_k3:
        hasil["reviu"]["E30"]["nilai"] = "Petugas K3 Konstruksi"
        hasil["reviu"]["E30"]["status"] = "terisi"
        hasil["reviu"]["E31"]["nilai"] = "-"
        hasil["reviu"]["E31"]["status"] = "terisi"
        sert_k3 = p_k3["sertifikat"]
        hasil["reviu"]["E32"]["nilai"] = sert_k3
        hasil["reviu"]["E32"]["status"] = "terisi" if sert_k3 else "tidak_ada"
    return True

def _parse_rk3k(pdf_or_path, hasil):
    """E33-E34: RK3K"""
    bahaya_list = []
    risiko_tertinggi_bahaya = ""
    risiko_tertinggi_uraian = ""
    
    if isinstance(pdf_or_path, str):
        if pdf_or_path.lower().endswith(".docx"):
            tables = extract_tables_from_docx(pdf_or_path)
            teks_full = extract_text_from_docx(pdf_or_path)
        else: return False
    else:
        tables = []
        for page in pdf_or_path.pages:
            t = page.extract_tables()
            if t: tables.extend(t)
        teks_full = "\n".join([p.extract_text() or "" for p in pdf_or_path.pages])

    for table in tables:
        if not table or len(table) < 2: continue
        idx_uraian = 1; idx_detail = 2; idx_bahaya = 3; idx_risiko = -1
        header = [str(c or "").lower() for c in table[0]]
        for ii, col in enumerate(header):
            if "identifikasi" in col and "bahaya" in col: idx_bahaya = ii
            if "uraian" in col and "pekerjaan" in col: idx_uraian = ii
            if "detail" in col or "item" in col: idx_detail = ii
            # Prioritas: sasaran k3 (berisi "Fatal/Cukup Fatal") > keterangan/tingkat/level > risiko umum
            if "sasaran" in col and "k3" in col: idx_risiko = ii; break
            elif any(kw in col for kw in ["keterangan", "tingkat", "level"]) and idx_risiko == -1: idx_risiko = ii
            elif "risiko" in col and idx_risiko == -1: idx_risiko = ii
        if idx_detail == -1 and idx_uraian != -1 and len(header) > idx_uraian + 1: idx_detail = idx_uraian + 1

        # Track baris dengan risiko tertinggi (prioritas: "fatal" tanpa "cukup")
        highest_score = 0
        for row in table[1:]:
            if len(row) > idx_bahaya:
                b_val = bersihkan(row[idx_bahaya])
                if b_val and len(b_val) > 5 and not b_val.isdigit() and "identifikasi" not in b_val.lower():
                    bahaya_list.append(b_val)
                    is_highest = False
                    score = 0
                    if idx_risiko != -1 and len(row) > idx_risiko:
                        r_val = str(row[idx_risiko] or "").lower()
                        if any(kw in r_val for kw in ["sangat fatal", "sangat besar", "ekstrem", "extreme", "meninggal"]): score = 4
                        elif re.search(r"\bfatal\b", r_val) and "cukup fatal" not in r_val: score = 3
                        elif any(kw in r_val for kw in ["tertinggi", "kritis", "high", "major", "cacat"]): score = 2
                        elif any(kw in r_val for kw in ["besar", "tinggi", "cukup fatal"]): score = 1
                        if score > highest_score:
                            highest_score = score
                            is_highest = True
                    if is_highest:
                        u_val = bersihkan(row[idx_uraian]) if idx_uraian != -1 and len(row) > idx_uraian else ""
                        d_val = bersihkan(row[idx_detail]) if idx_detail != -1 and len(row) > idx_detail else ""
                        risiko_tertinggi_uraian = f"{u_val} ({d_val})" if u_val and d_val else (u_val or d_val)
                        risiko_tertinggi_bahaya = b_val

    if not risiko_tertinggi_bahaya:
        m_rt = re.search(r"(?:Resiko|Tingkat\s+Risiko)\s+Tertinggi\s*[:\s]*([A-Z][^\n]{10,})", teks_full, re.IGNORECASE)
        if m_rt: risiko_tertinggi_bahaya = bersihkan(m_rt.group(1))

    # Fallback: pola PDF terpecah — "- <bahaya> Resiko\n<...> Paling\nTinggi"
    if not risiko_tertinggi_bahaya:
        m_pt = re.search(r"Paling\s*\nTinggi", teks_full, re.IGNORECASE)
        if m_pt:
            pre = teks_full[max(0, m_pt.start()-400):m_pt.start()]
            # Cari "- <bahaya> Resiko\n" di teks sebelum "Paling\nTinggi"
            m_pre = re.search(r"- ([A-Za-z][^\n]{3,60}) Resiko\n", pre, re.IGNORECASE)
            if m_pre:
                risiko_tertinggi_bahaya = bersihkan(m_pre.group(1))

    if hasil["reviu"]["E33"]["status"] != "terisi":
        if risiko_tertinggi_uraian:
            hasil["reviu"]["E33"]["nilai"] = risiko_tertinggi_uraian
            hasil["reviu"]["E33"]["status"] = "terisi"
        elif bahaya_list:
            hasil["reviu"]["E33"]["nilai"] = "\n".join(list(dict.fromkeys(bahaya_list))[:5])
            hasil["reviu"]["E33"]["status"] = "terisi"
    if hasil["reviu"]["E34"]["status"] != "terisi" and risiko_tertinggi_bahaya:
        hasil["reviu"]["E34"]["nilai"] = risiko_tertinggi_bahaya
        hasil["reviu"]["E34"]["status"] = "terisi"
    return len(bahaya_list) > 0

def _parse_dokpil_uraian(pdf_or_path, hasil):
    """E6-E15 Dokpil: Uraian Pekerjaan (Divisi)"""
    divisi_list = []
    if isinstance(pdf_or_path, str):
        if pdf_or_path.lower().endswith((".xlsx", ".xlsm")):
            divisi_list = extract_rekap_from_xlsx(pdf_or_path)
        elif pdf_or_path.lower().endswith(".docx"):
            teks_list = [extract_text_from_docx(pdf_or_path)]
        else: return False
    else:
        teks_list = [p.extract_text() or "" for p in pdf_or_path.pages]

    if not divisi_list:
        for teks in teks_list:
            matches = re.findall(r"^[ \t]*([IVXLC]+\.?)[ \t]+([A-Z][A-Z \/,()&]{5,100})", teks, re.MULTILINE)
            for m in matches:
                rom, desc = m
                desc = bersihkan(desc)
                if desc and desc not in divisi_list and len(desc) > 3:
                    if any(kw in desc.lower() for kw in ["pekerjaan", "divisi", "umum", "galian", "tanah", "aspal", "beton", "persiapan"]):
                        divisi_list.append(desc)
            if len(divisi_list) >= 2: break
        
    DOKPIL_CELLS = ["E6","E7","E8","E9","E10","E11","E12","E13","E14","E15"]
    for i, div in enumerate(divisi_list[:10]):
        hasil["dokpil"][DOKPIL_CELLS[i]]["nilai"] = div
        hasil["dokpil"][DOKPIL_CELLS[i]]["status"] = "terisi"
    for i in range(len(divisi_list), 10):
        hasil["dokpil"][DOKPIL_CELLS[i]]["status"] = "tidak_ada"
    return len(divisi_list) > 0

# ─── Parser Utama ────────────────────────────────────────────────────────────

def parse_pdf_enhanced(path_or_dir, bidang=""):
    try:
        import pdfplumber
    except ImportError: return None, "pdfplumber tidak tersedia"

    hasil = {
        "input_data": {
            "E16": {"nilai": "", "status": "kosong", "label": "Kegiatan/Sub Kegiatan", "sumber": ""},
            "E32": {"nilai": "", "status": "kosong", "label": "Lokasi Pekerjaan", "sumber": ""},
            "E33": {"nilai": "", "status": "kosong", "label": "Sumber Dana", "sumber": ""},
        },
        "reviu": {
            "E2":  {"nilai": "", "status": "kosong", "label": "Fungsi Bangunan", "sumber": ""},
            "E6":  {"nilai": "", "status": "kosong", "label": "SBU KBLI 2020", "sumber": ""},
            "E7":  {"nilai": "", "status": "kosong", "label": "SBU KBLI 2015", "sumber": ""},
            "E9":  {"nilai": "", "status": "kosong", "label": "Alat 1", "sumber": ""},
            "E10": {"nilai": "", "status": "kosong", "label": "Alat 2", "sumber": ""},
            "E11": {"nilai": "", "status": "kosong", "label": "Alat 3", "sumber": ""},
            "E12": {"nilai": "", "status": "kosong", "label": "Alat 4", "sumber": ""},
            "E13": {"nilai": "", "status": "kosong", "label": "Alat 5", "sumber": ""},
            "E14": {"nilai": "", "status": "kosong", "label": "Alat 6", "sumber": ""},
            "E15": {"nilai": "", "status": "kosong", "label": "Kapasitas 1", "sumber": ""},
            "E16": {"nilai": "", "status": "kosong", "label": "Kapasitas 2", "sumber": ""},
            "E17": {"nilai": "", "status": "kosong", "label": "Kapasitas 3", "sumber": ""},
            "E18": {"nilai": "", "status": "kosong", "label": "Kapasitas 4", "sumber": ""},
            "E19": {"nilai": "", "status": "kosong", "label": "Kapasitas 5", "sumber": ""},
            "E20": {"nilai": "", "status": "kosong", "label": "Kapasitas 6", "sumber": ""},
            "E21": {"nilai": "", "status": "kosong", "label": "Jumlah Alat 1", "sumber": ""},
            "E22": {"nilai": "", "status": "kosong", "label": "Jumlah Alat 2", "sumber": ""},
            "E23": {"nilai": "", "status": "kosong", "label": "Jumlah Alat 3", "sumber": ""},
            "E24": {"nilai": "", "status": "kosong", "label": "Jumlah Alat 4", "sumber": ""},
            "E25": {"nilai": "", "status": "kosong", "label": "Jumlah Alat 5", "sumber": ""},
            "E26": {"nilai": "", "status": "kosong", "label": "Jumlah Alat 6", "sumber": ""},
            "E27": {"nilai": "", "status": "kosong", "label": "Jabatan Teknis", "sumber": ""},
            "E28": {"nilai": 0,  "status": "kosong", "label": "Pengalaman Teknis (Tahun)", "sumber": ""},
            "E29": {"nilai": "", "status": "kosong", "label": "Nama SKK/SKT", "sumber": ""},
            "E30": {"nilai": "", "status": "kosong", "label": "Jabatan K3", "sumber": ""},
            "E31": {"nilai": 0,  "status": "kosong", "label": "Pengalaman K3 (Tahun)", "sumber": ""},
            "E32": {"nilai": "", "status": "kosong", "label": "Sertifikat K3", "sumber": ""},
            "E33": {"nilai": "", "status": "kosong", "label": "Uraian Pekerjaan RK3", "sumber": ""},
            "E34": {"nilai": "", "status": "kosong", "label": "Risiko Tertinggi RK3", "sumber": ""},
        },
        "dokpil": {
            "E6":  {"nilai": "", "status": "kosong", "label": "Uraian Pekerjaan 1", "sumber": ""},
            "E7":  {"nilai": "", "status": "kosong", "label": "Uraian Pekerjaan 2", "sumber": ""},
            "E8":  {"nilai": "", "status": "kosong", "label": "Uraian Pekerjaan 3", "sumber": ""},
            "E9":  {"nilai": "", "status": "kosong", "label": "Uraian Pekerjaan 4", "sumber": ""},
            "E10": {"nilai": "", "status": "kosong", "label": "Uraian Pekerjaan 5", "sumber": ""},
            "E11": {"nilai": "", "status": "kosong", "label": "Uraian Pekerjaan 6", "sumber": ""},
            "E12": {"nilai": "", "status": "kosong", "label": "Uraian Pekerjaan 7", "sumber": ""},
            "E13": {"nilai": "", "status": "kosong", "label": "Uraian Pekerjaan 8", "sumber": ""},
            "E14": {"nilai": "", "status": "kosong", "label": "Uraian Pekerjaan 9", "sumber": ""},
            "E15": {"nilai": "", "status": "kosong", "label": "Uraian Pekerjaan 10", "sumber": ""},
            "E16": {"nilai": "", "status": "kosong", "label": "Cara Pembayaran", "id_cp": "", "sumber": ""},
        },
        "error": ""
    }

    base_dir = path_or_dir if os.path.isdir(path_or_dir) else os.path.dirname(path_or_dir)
    # Draft_Pokja berada di subfolder 0. Draft Dokumen PPK; naik ke root paket
    # agar parser dapat memakai PDF Peralatan/Personil/Kuantitas yang tersedia.
    if os.path.basename(os.path.normpath(base_dir)).lower() == "0. draft dokumen ppk":
        base_dir = os.path.dirname(base_dir)
    file_map = {
        "alat": find_files_by_keywords(base_dir, ["Peralatan", "Alat"]),
        "personil": find_files_by_keywords(base_dir, ["Personil", "Tenaga"]),
        "rk3k": find_files_by_keywords(base_dir, ["RK3K", "SMKK", "Keselamatan", "K3"]),
        "dokpil": find_files_by_keywords(base_dir, ["Kuantitas", "BoQ", "RAB"]),
        "uraian": find_files_by_keywords(base_dir, ["Uraian Singkat", "KAK", "Spek"]),
        "merged": [path_or_dir] if os.path.isfile(path_or_dir) else find_files_by_keywords(base_dir, ["Draft_Pokja"]),
        "rks": find_files_by_keywords(base_dir, ["RKS", "Spesifikasi", "Teknis"]),
    }

    # Cache teks PDF per path — hindari re-open file yang sama berkali-kali
    _text_cache = {}
    _text_cache_lock = __import__("threading").Lock()

    def _get_pdf_text(path, max_pages=10):
        with _text_cache_lock:
            if path in _text_cache:
                return _text_cache[path]
        try:
            with pdfplumber.open(path) as _pdf:
                txt = "\n".join([p.extract_text() or "" for p in _pdf.pages[:max_pages]])
        except Exception:
            txt = ""
        with _text_cache_lock:
            _text_cache[path] = txt
        return txt

    # Helper untuk memproses file berdasarkan ekstensi
    def process_file(file_list, func, hasil):
        for f in file_list:
            ext = f.lower().split(".")[-1]
            if ext == "pdf":
                with pdfplumber.open(f) as pdf:
                    if func(pdf, hasil): return True
            elif ext in ["docx", "xlsx", "xlsm"]:
                if func(f, hasil): return True
        return False

    def _t1():
        process_file(file_map["uraian"] or file_map["rks"] or file_map["merged"], _parse_fung_bangunan, hasil)
    def _t2():
        process_file(file_map["alat"] or file_map["rks"] or file_map["merged"], _parse_peralatan, hasil)
    def _t3():
        process_file(file_map["personil"] or file_map["rks"] or file_map["merged"], _parse_personil, hasil)
    def _t4():
        # RK3K: hanya dari file yang namanya eksplisit mengandung RK3K/K3/Keselamatan
        # JANGAN parse dari Draft_Pokja — tabel di Draft PDF tidak reliabel untuk K3
        for f in file_map["rk3k"]:
            ext = f.lower().split(".")[-1]
            if ext == "pdf":
                with pdfplumber.open(f) as pdf:
                    _parse_rk3k(pdf, hasil)
            elif ext == "docx":
                _parse_rk3k(f, hasil)
    def _t5():
        process_file(file_map["dokpil"] or file_map["merged"], _parse_dokpil_uraian, hasil)

    # pdfplumber/PDF object tidak stabil saat beberapa file dibaca paralel
    # pada workbook paket; serial lebih lambat sedikit tetapi mencegah hang.
    with _cf.ThreadPoolExecutor(max_workers=1) as _pool:
        for _fut in _cf.as_completed([_pool.submit(t) for t in [_t1, _t2, _t3, _t4, _t5]]):
            _fut.result()

    # 3. Common Fields (Input Data) dari SEMUA file (biasanya ada di RK3K atau KAK)
    # Gunakan _get_pdf_text() — cache hasil extract agar file tidak dibuka ulang
    all_files = file_map["uraian"] + file_map["rk3k"] + file_map["personil"]
    if not all_files:
        all_files = file_map["merged"]
    for f in all_files:
        ext = f.lower().split(".")[-1]
        if ext == "pdf":
            txt = _get_pdf_text(f, max_pages=10)
        elif ext == "docx":
            txt = extract_text_from_docx(f)
        else: continue
        
        # E16: Kegiatan/Sub Kegiatan
        if not hasil["input_data"]["E16"]["nilai"]:
            _BUKAN_KEGIATAN = [
                "tujuan", "latar", "bermaksud", "menjalin", "melaksanakan", "dengan ini",
                "pemerintah", "kami", "bahwa", "dalam rangka", "untuk menjalin",
                "satuan kerja", "kuasa pengguna", "pengguna anggaran", "pejabat pembuat",
                "konsultan", "kontraktor", "pengawas", "seksama", "dokumen tersebut", "memahami benar",
            ]

            # Pola 1: judul KAK/RKS — PRIORITAS UTAMA karena paling bersih
            # Contoh: "KERANGKA ACUAN KERJA\nPEMBANGUAN BANGUNAN CYTOTOXIC - RSUD DATU SANGGUL"
            m_kak = re.search(
                r"(?:KERANGKA ACUAN KERJA|SPESIFIKASI TEKNIS|RENCANA KERJA DAN SYARAT)\s*\n+([^\n]{10,150})",
                txt, re.IGNORECASE
            )
            if m_kak:
                val = bersihkan(m_kak.group(1))
                val_lower = val.lower()
                if (len(val) > 10
                        and not any(kw in val_lower for kw in _BUKAN_KEGIATAN)
                        and not val_lower.startswith(("tahun", "nomor", "tentang", "oleh", "a.", "b.", "1."))):
                    hasil["input_data"]["E16"]["nilai"] = val
                    hasil["input_data"]["E16"]["status"] = "terisi"

            # Pola 2: Label eksplisit "Sub Kegiatan :" / "Kegiatan :" — hanya jika bukan "Satuan Kerja"
            if not hasil["input_data"]["E16"]["nilai"]:
                m_keg = re.search(
                    r"\bSub\s+Kegiatan\s*[:/]\s*([^\n]{5,150})",
                    txt, re.IGNORECASE
                )
                if m_keg:
                    val = bersihkan(m_keg.group(1))
                    val_lower = val.lower()
                    if (len(val) > 5
                            and not any(kw in val_lower for kw in _BUKAN_KEGIATAN)
                            and not val_lower.startswith(("dan ", "atau ", "ini ", "itu "))):
                        hasil["input_data"]["E16"]["nilai"] = val
                        hasil["input_data"]["E16"]["status"] = "terisi"

            # Pola 3: "Nama Paket :" — alternatif
            if not hasil["input_data"]["E16"]["nilai"]:
                m_paket = re.search(r"Nama\s+Paket\s*[:/]\s*([^\n]{5,150})", txt, re.IGNORECASE)
                if m_paket:
                    val = bersihkan(m_paket.group(1))
                    if len(val) > 5 and not any(kw in val.lower() for kw in _BUKAN_KEGIATAN):
                        hasil["input_data"]["E16"]["nilai"] = val
                        hasil["input_data"]["E16"]["status"] = "terisi"

            # Pola 4: fallback dari nama_tender (sudah dikirim via argfile baris ke-4)
            if not hasil["input_data"]["E16"]["nilai"]:
                nama_tender_fb = getattr(parse_pdf_enhanced, "_nama_tender_fallback", "")
                if nama_tender_fb:
                    hasil["input_data"]["E16"]["nilai"] = bersihkan(nama_tender_fb)
                    hasil["input_data"]["E16"]["status"] = "terisi"

        # E32: Lokasi Pekerjaan — fixed "Kabupaten Tapin" (tidak perlu parse)
        hasil["input_data"]["E32"]["nilai"] = "Kabupaten Tapin"
        hasil["input_data"]["E32"]["status"] = "terisi"

        # E32: Lokasi Pekerjaan (kode lama dipertahankan sebagai fallback jika suatu saat perlu)
        _BUKAN_LOKASI = [
            "resiko", "risiko", "tingkat", "uraian", "bahaya", "pengendalian",
            "ibprp", "rk3k", "k3", "pekerjaan persiapan", "fasilitas penunjang",
            "ruang lingkup", "pengadaan pekerjaan", "konsultan", "pemberi tugas",
        ]
        _KATA_LOKASI = [
            "kabupaten", "kota", "provinsi", "kecamatan", "kelurahan", "desa",
            "jalan", "jl.", "rsud", "puskesmas", "sdn", "dinas", "kantor",
            "gedung", "bangunan", "alamat", "lingkungan", "kawasan"
        ]
        if not hasil["input_data"]["E32"]["nilai"]:
            # Pola 1: nilai di baris berikutnya setelah label "Lokasi Pekerjaan"
            m_lok2 = re.search(r"Lokasi\s*(?:Pekerjaan)?\s*[:/]?\s*\n([^\n]{5,120})", txt, re.IGNORECASE)
            # Pola 2: nilai di baris yang sama setelah label
            m_lok1 = re.search(r"Lokasi\s*(?:Pekerjaan)?\s*[:/]\s*([^\n]{5,120})", txt, re.IGNORECASE)
            for m_try in [m_lok2, m_lok1]:
                if not m_try: continue
                val_lok = bersihkan(m_try.group(1))
                val_lower = val_lok.lower()
                if not val_lok: continue
                if any(kw in val_lower for kw in _BUKAN_LOKASI): continue
                # Nilai lokasi harus mengandung setidaknya satu kata penanda lokasi
                if not any(kw in val_lower for kw in _KATA_LOKASI) and len(val_lok) < 15: continue
                hasil["input_data"]["E32"]["nilai"] = val_lok
                hasil["input_data"]["E32"]["status"] = "terisi"
                break
            # Pola 3 (fallback): ambil dari judul dokumen — baris yang mengandung kata penanda lokasi
            # Prioritaskan baris yang HANYA berisi nama instansi/lokasi (bukan judul pekerjaan)
            if not hasil["input_data"]["E32"]["nilai"]:
                _JUDUL_PEKERJAAN = ["pembangunan", "rehabilitasi", "renovasi", "peningkatan",
                                    "pemeliharaan", "pengadaan", "belanja modal"]
                # Kata lokasi yang TIDAK termasuk kata pekerjaan (hapus "bangunan", "gedung", dll)
                _KATA_LOKASI_STRICT = [
                    "kabupaten", "kota", "provinsi", "kecamatan", "kelurahan", "desa",
                    "jalan", "jl.", "rsud", "puskesmas", "sdn", "smpn", "sman", "smkn",
                    "dinas", "kantor", "alamat", "lingkungan", "kawasan"
                ]
                kandidat_lokasi = []
                for line in txt.split("\n")[:25]:
                    line_clean = bersihkan(line)
                    line_lower = line_clean.lower()
                    if not (len(line_clean) > 5 and len(line_clean) < 100): continue
                    if not any(kw in line_lower for kw in _KATA_LOKASI_STRICT): continue
                    if any(kw in line_lower for kw in _BUKAN_LOKASI): continue
                    if any(kw in line_lower for kw in ["kerangka acuan", "rencana kerja", "spesifikasi", "kontrak"]): continue
                    has_pekerjaan = any(kw in line_lower for kw in _JUDUL_PEKERJAAN)
                    kandidat_lokasi.append((0 if has_pekerjaan else 1, line_clean))
                if kandidat_lokasi:
                    kandidat_lokasi.sort(key=lambda x: x[0], reverse=True)
                    hasil["input_data"]["E32"]["nilai"] = kandidat_lokasi[0][1]
                    hasil["input_data"]["E32"]["status"] = "terisi"

        # E33: Sumber Dana — semi-fix: deteksi APBD/APBDP dari PDF, sisa format difix
        # Format output: "APBD/APBDP Kabupaten Tapin Tahun Anggaran {tahun}"
        _tahun_berjalan = str(__import__("datetime").date.today().year)
        _jenis_apbd = "APBD"  # default
        txt_oneline_sd = re.sub(r"\s*\n\s*", " ", txt)
        m_jenis = re.search(r"\b(APBDP|APBD\s*[- ]?\s*P|APBD\s+PERUBAHAN|APBD)\b", txt_oneline_sd, re.IGNORECASE)
        if m_jenis:
            raw = m_jenis.group(1).upper()
            _jenis_apbd = "APBDP" if ("PERUBAHAN" in raw or "APBDP" in raw or re.search(r"APBD\s*[- ]\s*P\b", raw)) else "APBD"
        hasil["input_data"]["E33"]["nilai"] = f"{_jenis_apbd} Kabupaten Tapin Tahun Anggaran {_tahun_berjalan}"
        hasil["input_data"]["E33"]["status"] = "terisi"

    # 4. Cara Pembayaran
    daftar_cp = fetch_cara_pembayaran()
    id_cp, teks_cp = deteksi_cara_pembayaran(daftar_cp, bidang)
    hasil["dokpil"]["E16"]["nilai"] = teks_cp
    hasil["dokpil"]["E16"]["id_cp"] = id_cp
    hasil["dokpil"]["E16"]["status"] = "terisi"

    # 4b. Fallback E2 (Fungsi Bangunan) dari nama_tender jika masih kosong
    if hasil["reviu"]["E2"]["status"] != "terisi":
        nama_tender_arg = getattr(parse_pdf_enhanced, "_nama_tender_fallback", "")
        if nama_tender_arg:
            hasil["reviu"]["E2"]["nilai"] = bersihkan(nama_tender_arg)
            hasil["reviu"]["E2"]["status"] = "terisi"

    # 5. Normalisasi akhir: field yang masih "kosong" setelah semua proses → "tidak_ada"
    # (artinya sudah dicoba parse tapi tidak ditemukan)
    nama_tender_arg = getattr(parse_pdf_enhanced, "_nama_tender_fallback", "")
    if hasil["reviu"]["E6"]["status"] != "terisi" and nama_tender_arg:
        sbu_baru, sbu_lama = _fallback_sbu_dari_nama(nama_tender_arg)
        if sbu_baru:
            hasil["reviu"]["E6"]["nilai"] = sbu_baru
            hasil["reviu"]["E6"]["status"] = "terisi"
            hasil["reviu"]["E7"]["nilai"] = sbu_lama
            hasil["reviu"]["E7"]["status"] = "terisi"

    FIELD_PARSE_ONLY = [
        ("reviu", "E6"), ("reviu", "E7"),                          # SBU (dari PDF)
        ("reviu", "E27"), ("reviu", "E28"), ("reviu", "E29"),      # Personil teknis
        ("reviu", "E30"), ("reviu", "E31"), ("reviu", "E32"),      # Personil K3
        ("reviu", "E33"), ("reviu", "E34"),                        # RK3K
    ]
    for bagian, k in FIELD_PARSE_ONLY:
        if hasil[bagian][k]["status"] == "kosong":
            hasil[bagian][k]["status"] = "tidak_ada"
            # E28/E31 (pengalaman): jika nilai 0 dan tidak ada personil, kosongkan
            if k in ("E28", "E31") and hasil[bagian][k]["nilai"] == 0:
                hasil[bagian][k]["nilai"] = ""

    # Backfill sumber dengan nama file input
    _src_name = os.path.basename(path_or_dir)
    for sec in ("input_data", "reviu", "dokpil"):
        for key, item in hasil[sec].items():
            if isinstance(item, dict) and item.get("status") not in ("kosong", "") and item.get("sumber") == "":
                item["sumber"] = _src_name

    return hasil, ""

def main():
    bidang = ""
    nama_tender = ""
    if len(sys.argv) >= 3 and sys.argv[1] == "--argfile":
        argfile = sys.argv[2]
        with open(argfile, encoding="utf-8") as f:
            lines = [l.rstrip("\n").rstrip("\r") for l in f.readlines()]
        path_pdf = lines[0] if len(lines) > 0 else ""
        folder_output = lines[1] if len(lines) > 1 else ""
        bidang = lines[2] if len(lines) > 2 else ""
        nama_tender = lines[3] if len(lines) > 3 else ""
    elif len(sys.argv) >= 3:
        path_pdf = sys.argv[1]
        folder_output = sys.argv[2]
        bidang = sys.argv[3] if len(sys.argv) > 3 else ""
        nama_tender = sys.argv[4] if len(sys.argv) > 4 else ""
    else:
        print("Usage: python parse_reviu.py <path_pdf_or_dir> <folder_output> [bidang] [nama_tender]")
        sys.exit(1)

    # Kirim nama_tender sebagai fallback untuk E2 via attribute sementara
    parse_pdf_enhanced._nama_tender_fallback = nama_tender
    hasil, err = parse_pdf_enhanced(path_pdf, bidang)
    if err:
        hasil = {"error": err, "reviu": {}, "dokpil": {}, "input_data": {}}

    out_path = os.path.join(folder_output, "_parse_reviu.json")
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(hasil, f, ensure_ascii=False, indent=2)
    print(f"OK: {out_path}")

if __name__ == "__main__":
    main()
