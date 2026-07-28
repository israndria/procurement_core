"""Patch legacy Ringkasan Evaluasi DOCX while preserving the package layout."""
from __future__ import annotations

import re
import sys
from docx import Document


def _set_cell_text(cell, text: str) -> None:
    # Keep first run formatting; clear extra paragraphs/runs only when needed.
    if not cell.paragraphs:
        cell.text = text
        return
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    for extra in cell.paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


def sanitize(path: str) -> int:
    doc = Document(path)
    removed_rows = 0
    for table in doc.tables:
        for row in list(table.rows):
            texts = [cell.text for cell in row.cells]
            joined = "\n".join(texts)
            if "KBLI 2017" in joined:
                row._tr.getparent().remove(row._tr)
                removed_rows += 1
                continue
            for cell in row.cells:
                text = cell.text
                if "c)." in text and "verifikasi" in text:
                    text = re.sub(
                        r"b\)\.(.*?verifikasi)\s*;?\s*atau\s*c\)\..*",
                        lambda m: "b)." + m.group(1).strip() + ".",
                        text,
                        flags=re.I | re.S,
                    )
                text = re.sub(r"\s*\(\s*\)", "", text)
                if text != cell.text:
                    _set_cell_text(cell, text)
    doc.save(path)
    return removed_rows


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("usage: sanitize_revaluasi_docx.py FILE.docx")
    print(f"removed_rows={sanitize(sys.argv[1])}")
